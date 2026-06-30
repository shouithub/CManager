# type: ignore[attr-defined]
"""
Django 社团管理系统视图模块

此文件包含大量 Django 模型交互代码。由于 Pylance 无法完全识别 Django ORM 的动态特性
（如自动生成的 id 字段、相关管理器等），我们在文件顶部添加全局类型忽略指令。
这对代码的实际功能没有影响，只是消除了 IDE 中的假性错误警告。
"""
from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse
from django.views.decorators.http import require_http_methods, require_GET, require_POST
from django.contrib.auth.decorators import login_required
from django.contrib.auth import authenticate
from django.contrib.auth.models import User
from django.contrib import messages
from django.utils import timezone
from datetime import datetime
from decimal import Decimal, InvalidOperation
from django.conf import settings
from django.http import HttpResponse, FileResponse, HttpResponseForbidden, JsonResponse, Http404
from django.db import IntegrityError, transaction
from django.db.models import Q, Prefetch, FileField, Count
from django.core.cache import cache
from django.core.files.base import ContentFile
from collections import defaultdict
import os
import re
import urllib.parse
import os
import base64
import tempfile
import csv
import io
import json
from .models import Club, Officer, UserProfile, FormChannel, FormCycle, FormChannelClubState, FormField, FormSubmission, FormSubmissionReview, FormFieldValue, FormUploadedFile, Template, Announcement, StaffClubRelation, SMTPConfig, CarouselImage, Department, Room, RoomBooking, TimeSlot, SiteSettings, DailyStat, ClubMember, RegistrationToken, PublishedActivity, ActivityRegistration
from .business_forms import (
    BusinessActionError,
    apply_business_action,
    create_form_cycle,
    is_locked_business_field,
    locked_business_field_keys,
    missing_required_field_keys,
    seed_business_form_channels,
)
from django.contrib.contenttypes.models import ContentType
import shutil
from PIL import Image
from .context_processors import audit_center_counts as get_audit_center_counts
from .site_assets import process_site_logo
from .lifecycle_utils import mark_profile_inactive


def rename_uploaded_file(file, club_name, request_type, material_type):
    """
    为上传的文件重命名为：社团名-请求类型-文件类型
    例如：社团名-年审-自查表.docx

    Args:
        file: 上传的文件对象
        club_name: 社团名称
        request_type: 请求类型（'年审', '报销', '注册'等）
        material_type: 文件类型（'自查表', '报销凭证'等）

    Returns:
        修改后的文件对象
    """
    if not file:
        return file

    # 获取文件扩展名
    file_ext = os.path.splitext(file.name)[1]

    # 生成新的文件名
    new_filename = f"{club_name}-{request_type}-{material_type}{file_ext}"

    # 清理特殊字符，避免文件系统问题
    invalid_chars = '<>:"/\\|?*'
    for char in invalid_chars:
        new_filename = new_filename.replace(char, '_')

    # 修改文件名
    file.name = new_filename
    return file


def _is_president(user):
    """检查用户是否为社长"""
    if not user.is_authenticated:
        return False
    try:
        return user.profile.role == 'president'
    except UserProfile.DoesNotExist:
        return False


def _is_staff(user):
    """检查用户是否为干事"""
    if not user.is_authenticated:
        return False
    try:
        return user.profile.role == 'staff'
    except UserProfile.DoesNotExist:
        return False


def _is_admin(user):
    """检查用户是否为管理员"""
    if not user.is_authenticated:
        return False
    try:
        return user.profile.role == 'admin'
    except UserProfile.DoesNotExist:
        return False


def _get_president_club_ids(user):
    """获取社长可访问的社团ID（通过 Officer 表查询）。"""
    return list(Officer.objects.filter(
        user_profile__user=user,
        position='president',
        is_current=True,
    ).values_list('club_id', flat=True))


def _build_external_url(request, path: str) -> str:
    """Build a public URL for client-facing links behind reverse proxies."""
    forwarded_proto = (request.headers.get('X-Forwarded-Proto') or '').split(',')[0].strip()
    forwarded_host = (request.headers.get('X-Forwarded-Host') or '').split(',')[0].strip()

    scheme = forwarded_proto or request.scheme
    host = forwarded_host or request.get_host()

    # Some proxy setups still expose loopback host to Django; fallback to browser origin.
    loopback_hosts = {'127.0.0.1', 'localhost', '[::1]'}
    host_without_port = host.split(':', 1)[0].strip().lower()
    if host_without_port in loopback_hosts:
        for header_name in ('Origin', 'Referer'):
            header_value = request.headers.get(header_name, '').strip()
            if not header_value:
                continue
            parsed = urllib.parse.urlparse(header_value)
            if parsed.netloc:
                scheme = parsed.scheme or scheme
                host = parsed.netloc
                break

    return f'{scheme}://{host}{path}'


def _make_qr_data_uri(payload: str) -> str:
    """Generate a PNG QR code and return as data URI; return empty string on failure."""
    try:
        import qrcode  # type: ignore
        qr = qrcode.QRCode(version=1, box_size=8, border=2)
        qr.add_data(payload)
        qr.make(fit=True)
        img = qr.make_image(fill_color='black', back_color='white')
        buffer = io.BytesIO()
        img.save(buffer, format='PNG')
        return 'data:image/png;base64,' + base64.b64encode(buffer.getvalue()).decode('ascii')
    except Exception:
        return ''


def is_staff_or_admin(user):
    """返回用户是否为干事或管理员（布尔）。超级用户也视为管理员。"""
    try:
        # 超级用户始终有管理员权限
        if getattr(user, 'is_superuser', False):
            return True
        return getattr(user, 'profile', None) and user.profile.role in ['staff', 'admin']
    except Exception:
        return False


def _validate_word_file(file, field_name):
    """验证上传文件是否为 Word 格式（.doc/.docx）。

    返回: 错误消息字符串或 None
    """
    if not file:
        return f"{field_name} 文件不能为空"
    valid_extensions = ['.doc', '.docx']
    ext = os.path.splitext(file.name)[1].lower()
    if ext not in valid_extensions:
        return f"{field_name} 必须为 Word 文档 (.doc 或 .docx)"
    return None


def _validate_file_allowed(file, field_name, allowed_extensions, allowed_mimetypes=None):
    """通用文件类型验证函数。返回错误消息或 None。"""
    if not file:
        return f"{field_name} 文件不能为空"
    ext = os.path.splitext(file.name)[1].lower()
    if ext not in allowed_extensions:
        return f"{field_name} 的文件类型不被允许（允许的后缀：{', '.join(allowed_extensions)}）"
    # 可选的 mime 类型检查（如果需要）
    return None










import json

@login_required(login_url=settings.LOGIN_URL)
def download_file(request):
    """自定义文件下载视图，用于处理文件下载并重命名

    GET参数:
        file_path: 文件的相对路径（相对于MEDIA_ROOT）
        filename: 下载时使用的文件名
    """
    # 从GET请求中获取参数
    file_path = request.GET.get('file_path', '')
    filename = request.GET.get('filename', '')

    # 添加调试信息
    debug_info = {
        'received_params': {
            'file_path': file_path,
            'filename': filename,
        },
        'processing_steps': [],
        'settings_info': {
            'MEDIA_ROOT': str(getattr(settings, 'MEDIA_ROOT', 'Not set')),
            'MEDIA_URL': str(getattr(settings, 'MEDIA_URL', 'Not set')),
            'BASE_DIR': str(getattr(settings, 'BASE_DIR', 'Not set')),
        }
    }

    debug_info['processing_steps'].append(f"Received parameters - file_path: {file_path}, filename: {filename}")

    # 检查必要参数
    if not file_path:
        debug_info['processing_steps'].append('Missing file_path parameter')
        response = HttpResponse("缺少文件路径参数", status=400)
        response['X-Debug-Info'] = json.dumps(debug_info, ensure_ascii=False)
        return response

    # 构建完整的文件路径
    # 清理file_path，移除可能的查询参数或片段
    if '?' in file_path:
        file_path = file_path.split('?')[0]
        debug_info['processing_steps'].append(f"Removed query parameters: {file_path}")
    if '#' in file_path:
        file_path = file_path.split('#')[0]
        debug_info['processing_steps'].append(f"Removed fragment: {file_path}")

    # 如果file_path包含完整URL，提取相对路径
    if file_path.startswith('http://') or file_path.startswith('https://'):
        # 移除域名部分，获取相对路径
        from urllib.parse import urlparse
        parsed_url = urlparse(file_path)
        file_path = parsed_url.path
        debug_info['processing_steps'].append(f"Extracted path from URL: {file_path}")

        # 如果路径以MEDIA_URL开头，移除它
        media_url = settings.MEDIA_URL
        if file_path.startswith(media_url):
            file_path = file_path[len(media_url):]
            debug_info['processing_steps'].append(f"Removed MEDIA_URL prefix: {file_path}")
        elif file_path.startswith('/' + media_url):
            file_path = file_path[len('/' + media_url):]
            debug_info['processing_steps'].append(f"Removed leading slash and MEDIA_URL prefix: {file_path}")
    elif file_path.startswith('/'):
        # 如果路径以斜杠开头，移除它
        file_path = file_path[1:]
        debug_info['processing_steps'].append(f"Removed leading slash: {file_path}")

    # 特别处理以media/开头的路径
    if file_path.startswith('media/'):
        file_path = file_path[6:]  # 移除'media/'前缀
        debug_info['processing_steps'].append(f"Removed media/ prefix: {file_path}")

    full_path = os.path.join(settings.MEDIA_ROOT, file_path)
    debug_info['processing_steps'].append(f"Constructed full path: {full_path}")
    debug_info['processing_steps'].append(f"MEDIA_ROOT: {settings.MEDIA_ROOT}")

    # 检查文件是否存在
    if not os.path.exists(full_path):
        debug_info['processing_steps'].append(f"File not found at primary path: {full_path}")
        # 尝试其他可能的路径
        alternative_path = os.path.join(settings.BASE_DIR, file_path)
        debug_info['processing_steps'].append(f"Trying alternative path: {alternative_path}")
        if os.path.exists(alternative_path):
            full_path = alternative_path
            debug_info['processing_steps'].append(f"Found file at alternative path: {full_path}")
        else:
            debug_info['processing_steps'].append(f"Alternative path also not found: {alternative_path}")
            response = HttpResponse(json.dumps(debug_info, ensure_ascii=False, indent=2), content_type='application/json', status=404)
            response['X-Debug-Info'] = json.dumps(debug_info, ensure_ascii=False)
            return response

    debug_info['processing_steps'].append(f"File exists: {full_path}")

    # 获取文件大小
    try:
        file_size = os.path.getsize(full_path)
        debug_info['processing_steps'].append(f"File size: {file_size} bytes")
    except Exception as e:
        debug_info['processing_steps'].append(f"Error getting file size: {str(e)}")
        file_size = None

    # 如果没有提供文件名，使用原始文件名
    if not filename:
        filename = os.path.basename(full_path)
        debug_info['processing_steps'].append(f"Using default filename: {filename}")
    else:
        # 确保文件名是安全的（移除路径分隔符）
        filename = os.path.basename(filename)
        debug_info['processing_steps'].append(f"Using provided filename: {filename}")

    # 添加最终调试信息
    debug_info['processing_steps'].append(f"Final filename: {filename}")

    # 创建文件响应
    try:
        # 打开文件并创建响应
        with open(full_path, 'rb') as f:
            response = HttpResponse(f.read(), content_type='application/octet-stream')

            # 处理文件名编码以支持中文
            ascii_filename = filename.encode('ascii', 'ignore').decode('ascii')
            utf8_filename = filename.encode('utf-8')

            # 设置Content-Disposition头以指定下载文件名
            response['Content-Disposition'] = f"attachment; filename*=UTF-8''{urllib.parse.quote(utf8_filename)}"
            debug_info['processing_steps'].append(f"Set Content-Disposition: attachment; filename*=UTF-8''{urllib.parse.quote(utf8_filename)}")

            # 添加文件信息到响应头
            if file_size:
                response['Content-Length'] = str(file_size)
                debug_info['processing_steps'].append(f"Set Content-Length: {file_size}")

            # 添加调试信息到响应头
            response['X-Debug-Info'] = json.dumps(debug_info, ensure_ascii=False)

            return response

    except Exception as e:
        debug_info['processing_steps'].append(f"Error opening file: {str(e)}")
        response = HttpResponse(json.dumps(debug_info, ensure_ascii=False, indent=2), content_type='application/json', status=500)
        response['X-Debug-Info'] = json.dumps(debug_info, ensure_ascii=False)
        return response



@login_required(login_url='clubs:login')
def user_detail(request, user_id):
    """用户详情页 - 显示用户公开信息及关联社团/干事"""
    target_user = get_object_or_404(User, pk=user_id)
    from .models import StaffClubRelation, Officer

    context = {
        'target_user': target_user,
        'responsible_clubs': [],
        'affiliated_clubs': [],
        'affiliated_clubs_with_staff': [],
        'responsible_staff_list': [],
        'is_staff_of_viewing_president': False,
    }

    try:
        profile = target_user.profile

        # 检查是否为当前查看者(社长)负责的社团的干事
        if request.user.is_authenticated and hasattr(request.user, 'profile') and request.user.profile.role == 'president':
            if profile.role == 'staff':
                # 获取该干事负责的社团
                staff_club_ids = StaffClubRelation.objects.filter(
                    staff=profile,
                    is_active=True
                ).values_list('club_id', flat=True)

                # 获取当前用户(社长)负责的社团
                president_club_ids = Officer.objects.filter(
                    user_profile=request.user.profile,
                    position='president',
                    is_current=True
                ).values_list('club_id', flat=True)

                # 检查是否有交集
                if set(staff_club_ids) & set(president_club_ids):
                    context['is_staff_of_viewing_president'] = True

        # 如果是干事或管理员，获取负责的社团
        if profile.role in ['staff', 'admin']:
            context['responsible_clubs'] = StaffClubRelation.objects.filter(
                staff=profile,
                is_active=True
            ).select_related('club')

        # 如果是社长，获取所属社团及对应的负责干事
        if profile.role == 'president':
            # 获取当前担任职位的社团
            officer_positions = Officer.objects.filter(
                user_profile=profile,
                is_current=True
            ).select_related('club')
            context['affiliated_clubs'] = officer_positions

            # 获取这些社团对应的负责干事
            club_ids = officer_positions.values_list('club_id', flat=True)
            staff_relations = StaffClubRelation.objects.filter(
                club_id__in=club_ids,
                is_active=True
            ).select_related('staff', 'staff__user').distinct()

            staff_by_club = {}
            staff_seen_by_club = {}
            for relation in staff_relations:
                club_id = relation.club_id
                staff_by_club.setdefault(club_id, [])
                staff_seen_by_club.setdefault(club_id, set())
                if relation.staff_id in staff_seen_by_club[club_id]:
                    continue
                staff_by_club[club_id].append(relation.staff)
                staff_seen_by_club[club_id].add(relation.staff_id)

            context['affiliated_clubs_with_staff'] = [
                {
                    'officer': officer,
                    'club': officer.club,
                    'staff_list': staff_by_club.get(officer.club_id, []),
                }
                for officer in officer_positions
            ]

    except UserProfile.DoesNotExist:
        pass

    return render(request, 'clubs/user_detail.html', context)


def index(request):
    """首页 - 显示部门介绍、社团信息和最新公告"""
    from .models import Department

    # 共享数据（所有用户类型复用，缓存30秒）
    departments = cache.get('index:departments')
    if departments is None:
        departments = list(Department.objects.all().order_by('order'))
        cache.set('index:departments', departments, 30)

    announcements = cache.get('index:announcements')
    if announcements is None:
        announcements = list(Announcement.objects.filter(status='published').order_by('-published_at')[:5])
        cache.set('index:announcements', announcements, 30)

    carousel_images = cache.get('index:carousel_images')
    if carousel_images is None:
        carousel_images = list(CarouselImage.objects.filter(is_active=True).order_by('order', '-uploaded_at'))
        cache.set('index:carousel_images', carousel_images, 30)

    # 未登录用户显示部门介绍和公告
    if not request.user.is_authenticated:
        context = {
            'is_anonymous': True,
            'departments': departments,
            'announcements': announcements,
            'carousel_images': carousel_images,
        }
        return render(request, 'clubs/index.html', context)

    # 检查是否为干事或管理员
    staff_admin = is_staff_or_admin(request.user)

    if staff_admin:
        # 为干事和管理员显示部门介绍和树状图
        from .models import StaffClubRelation

        # 获取组织统计
        total_staff = UserProfile.objects.filter(role='staff', status='approved').count()
        total_directors = UserProfile.objects.filter(role='staff', staff_level='director').count()
        total_members = UserProfile.objects.filter(role='staff', staff_level='member').count()

        # 获取所有干事，消除 N+1：一次性批量查出所有干事-社团关系
        staff_users = list(
            UserProfile.objects.filter(role='staff', status='approved')
            .select_related('user')
            .only('id', 'real_name', 'department', 'staff_level', 'user__username')
        )
        staff_ids = [s.id for s in staff_users]

        relations_qs = (
            StaffClubRelation.objects
            .filter(staff_id__in=staff_ids, is_active=True)
            .select_related('club')
            .only('staff_id', 'club__id', 'club__name', 'club__status',
                  'club__members_count', 'club__founded_date', 'club__description')
        )
        clubs_by_staff: dict = defaultdict(list)
        for rel in relations_qs:
            clubs_by_staff[rel.staff_id].append({
                'id': rel.club.id,
                'name': rel.club.name,
                'status': rel.club.status,
                'members_count': rel.club.members_count,
                'founded_date': rel.club.founded_date,
                'description': rel.club.description,
            })

        staff_tree_data = []
        for staff_profile in staff_users:
            clubs = clubs_by_staff.get(staff_profile.id, [])
            staff_tree_data.append({
                'staff_id': staff_profile.id,
                'staff_name': staff_profile.get_full_name(),
                'staff_username': staff_profile.user.username,
                'clubs': clubs,
                'clubs_count': len(clubs),
            })

        context = {
            'is_staff_or_admin': staff_admin,
            'departments': departments,
            'total_staff': total_staff,
            'total_directors': total_directors,
            'total_members': total_members,
            'can_edit': request.user.profile.role == 'admin',
            'staff_tree_data': staff_tree_data,
            'announcements': announcements,
            'carousel_images': carousel_images,
            'total_clubs': Club.objects.count(),
        }
        return render(request, 'clubs/index.html', context)

    # 普通用户显示所有社团
    clubs = list(Club.objects.all())
    # 一次性查出当前用户担任社长的所有社团 ID，消除 N+1
    president_club_ids = set(
        Officer.objects.filter(
            user_profile__user=request.user,
            position='president',
            is_current=True
        ).values_list('club_id', flat=True)
    )

    clubs_data = [
        {'club': club, 'is_president': club.id in president_club_ids}
        for club in clubs
    ]

    context = {
        'clubs_data': clubs_data,
        'clubs': clubs,
        'announcements': announcements,
        'carousel_images': carousel_images,
        'departments': departments,
        'total_clubs': len(clubs),
    }
    return render(request, 'clubs/index.html', context)


def club_detail(request, club_id):
    """社团详情页"""
    club = get_object_or_404(Club, pk=club_id)
    officers = Officer.objects.filter(club=club, is_current=True)
    memberships = ClubMember.objects.filter(club=club).select_related('user_profile__user').order_by('-joined_at')

    # 检查当前用户是否为该社团的社长
    is_president = False
    is_staff = False
    if request.user.is_authenticated:
        # 检查是否为社长
        is_president = Officer.objects.filter(
            user_profile__user=request.user,
            club=club,
            position='president',
            is_current=True
        ).exists()

        # 检查是否为干事或管理员
        try:
            is_staff = request.user.profile.role in ['staff', 'admin']
        except:
            is_staff = False

    context = {
        'club': club,
        'officers': officers,
        'memberships': memberships,
        'is_president': is_president,
        'is_staff': is_staff,
    }
    return render(request, 'clubs/club_detail.html', context)


@login_required(login_url=settings.LOGIN_URL)
@require_http_methods(['POST'])
def generate_member_join_token(request, club_id):
    """社长生成社员入会二维码令牌，支持可配置有效期和使用次数。"""
    club = get_object_or_404(Club, pk=club_id)

    is_club_president = Officer.objects.filter(
        user_profile__user=request.user,
        club=club,
        position='president',
        is_current=True,
    ).exists()
    if not is_club_president:
        return JsonResponse({'success': False, 'message': '仅该社团社长可生成入会二维码'}, status=403)

    # 获取参数：有效时长（分钟）和最大使用次数
    try:
        minutes = int(request.POST.get('minutes', 10))
        max_uses_str = request.POST.get('max_uses', '1')

        # 验证有效时长范围
        if minutes < 1 or minutes > 43200:  # 最多30天
            return JsonResponse({'success': False, 'message': '有效时长必须在1-43200分钟之间'}, status=400)

        # 解析使用次数
        if max_uses_str == 'unlimited':
            max_uses = None
            # 业务规则：不限次数时，有效期最多1天
            if minutes > 1440:
                return JsonResponse({'success': False, 'message': '不限次数的令牌有效期不得超过1天（1440分钟）'}, status=400)
        else:
            max_uses = int(max_uses_str)
            if max_uses < 1:
                return JsonResponse({'success': False, 'message': '使用次数必须大于0'}, status=400)
    except (ValueError, TypeError):
        return JsonResponse({'success': False, 'message': '参数格式错误'}, status=400)

    token = RegistrationToken.create_for_club(club=club, created_by=request.user, minutes=minutes, max_uses=max_uses)
    join_path = reverse('clubs:member_join_by_token', args=[token.code])
    join_url = _build_external_url(request, join_path)
    qr_data_uri = _make_qr_data_uri(join_url)
    qr_error = '' if qr_data_uri else '二维码图片生成失败，请确认已安装 qrcode[pil] 依赖'

    uses_info = '不限次数' if max_uses is None else f'{max_uses}次'
    return JsonResponse({
        'success': True,
        'token': token.code,
        'expires_at': token.expires_at.strftime('%Y-%m-%d %H:%M:%S'),
        'join_url': join_url,
        'join_path': join_path,
        'qr_payload': join_url,
        'qr_data_uri': qr_data_uri,
        'qr_error': qr_error,
        'uses_info': uses_info,
    })


@login_required
@require_POST
def delete_member_token(request, club_id, token_id):
    """社长删除指定招新令牌。"""
    club = get_object_or_404(Club, pk=club_id)
    is_club_president = Officer.objects.filter(
        user_profile__user=request.user,
        club=club,
        position='president',
        is_current=True,
    ).exists()
    if not is_club_president:
        return JsonResponse({'success': False, 'message': '仅该社团社长可删除令牌'}, status=403)
    deleted, _ = RegistrationToken.objects.filter(pk=token_id, club=club).delete()
    if deleted:
        return JsonResponse({'success': True})
    return JsonResponse({'success': False, 'message': '令牌不存在'}, status=404)


@login_required
def list_member_tokens(request, club_id):
    """社长查看当前有效的招新令牌列表（GET JSON）。"""
    club = get_object_or_404(Club, pk=club_id)
    is_club_president = Officer.objects.filter(
        user_profile__user=request.user,
        club=club,
        position='president',
        is_current=True,
    ).exists()
    if not is_club_president:
        return JsonResponse({'success': False, 'message': '仅该社团社长可查看令牌'}, status=403)
    now = timezone.now()
    tokens = RegistrationToken.objects.filter(
        club=club,
        expires_at__gt=now,
    ).exclude(
        # 排除一次性且已使用完的令牌
        max_uses=1,
        used_count__gte=1,
    ).order_by('-created_at')
    data = []
    for t in tokens:
        join_path = reverse('clubs:member_join_by_token', args=[t.code])
        join_url = request.build_absolute_uri(join_path)
        uses_info = '不限次数' if t.max_uses is None else f'{t.used_count}/{t.max_uses}次'
        data.append({
            'id': t.id,
            'join_url': join_url,
            'join_path': join_path,
            'expires_at': t.expires_at.strftime('%Y-%m-%d %H:%M'),
            'uses_info': uses_info,
            'created_at': t.created_at.strftime('%Y-%m-%d %H:%M'),
        })
    return JsonResponse({'success': True, 'tokens': data})


@require_http_methods(['GET', 'POST'])
def member_join_by_token(request, token_code):
    """扫码加入社团：支持新建member账号或已有账号绑定，学号必填。"""
    token = get_object_or_404(RegistrationToken, code=token_code)

    if not token.can_use():
        messages.error(request, '注册链接已失效（已使用或已过期），请联系社长重新生成二维码')
        return redirect('clubs:index')

    if request.method == 'POST':
        existing_account = request.POST.get('existing_account', 'no') == 'yes'

        username = request.POST.get('username', '').strip()
        password = request.POST.get('password', '').strip()
        email = request.POST.get('email', '').strip()
        real_name = request.POST.get('real_name', '').strip()
        student_id = request.POST.get('student_id', '').strip()
        gender = request.POST.get('gender', '').strip()
        college = request.POST.get('college', '').strip()
        class_name = request.POST.get('class_name', '').strip()
        phone = request.POST.get('phone', '').strip()
        qq = request.POST.get('qq', '').strip()
        wechat = request.POST.get('wechat', '').strip()

        errors = []

        user = None
        if existing_account:
            if not username or not password:
                errors.append('已有账户绑定需要填写用户名和密码')
            else:
                user = authenticate(request, username=username, password=password)
                if user is None:
                    errors.append('用户名或密码不正确')
        else:
            if not real_name:
                errors.append('姓名不能为空')
            if not student_id:
                errors.append('学号不能为空')
            if not gender:
                errors.append('性别不能为空')
            if not college:
                errors.append('学院不能为空')
            if not class_name:
                errors.append('班级不能为空')
            if not phone:
                errors.append('手机号不能为空')
            if not wechat:
                errors.append('微信号不能为空')
            if not email:
                errors.append('邮箱不能为空')
            if not username:
                errors.append('用户名不能为空')
            if not password or len(password) < 6:
                errors.append('密码至少6位')
            if User.objects.filter(username=username).exists():
                errors.append('用户名已存在')
            if User.objects.filter(email=email).exists():
                errors.append('邮箱已被使用')

        if student_id and not existing_account:
            qs = UserProfile.objects.filter(student_id=student_id)
            if user is not None:
                qs = qs.exclude(user=user)
            if qs.exists():
                errors.append('学号已被使用')

        if errors:
            return render(request, 'clubs/member_join_form.html', {
                'token': token,
                'club': token.club,
                'errors': errors,
                'form_data': request.POST,
            })

        if user is None:
            user = User.objects.create_user(username=username, email=email, password=password, first_name=real_name)

        profile, _created = UserProfile.objects.get_or_create(
            user=user,
            defaults={
                'role': 'member',
                'status': 'approved',
                'account_status': 'active',
                'real_name': real_name,
                'student_id': student_id,
                'gender': gender,
                'college': college,
                'class_name': class_name,
                'phone': phone,
                'qq': qq,
                'wechat': wechat,
            },
        )

        if not existing_account:
            profile.role = 'member'
            profile.status = 'approved'
            profile.account_status = 'active'
            profile.real_name = real_name
            profile.student_id = student_id
            profile.gender = gender
            profile.college = college
            profile.class_name = class_name
            profile.phone = phone
            profile.qq = qq
            profile.wechat = wechat
            profile.save()

            user.email = email
            user.first_name = real_name
            user.save(update_fields=['email', 'first_name'])

        ClubMember.objects.get_or_create(
            club=token.club,
            user_profile=profile,
            defaults={'status': 'active'},
        )

        # 标记令牌已使用，智能处理使用次数
        token.mark_used()

        messages.success(request, f'已成功加入社团「{token.club.name}」，请使用账户登录系统')
        return redirect('clubs:login')

    return render(request, 'clubs/member_join_form.html', {
        'token': token,
        'club': token.club,
    })


















# ==================== 干事审核界面 ====================






# ==================== 报销功能 ====================





# ==================== 干事管理功能 ====================

@login_required(login_url=settings.LOGIN_URL)
def get_templates_by_type(template_type):
    """根据模板类型获取活跃的模板列表"""
    return Template.objects.filter(template_type=template_type, is_active=True).order_by('-created_at')








# ==================== 管理员功能 ====================

@login_required(login_url=settings.LOGIN_URL)
def review_staff_registration(request, user_id):
    """
    审核干事注册申请 - 仅管理员可用
    """
    if not _is_admin(request.user):
        messages.error(request, '仅管理员可以审核干事注册申请')
        return redirect('clubs:index')

    # 获取用户和用户角色信息
    user = get_object_or_404(User, pk=user_id)
    try:
        profile = user.profile
        # 确保只审核干事角色且状态为待审核的用户
        if profile.role != 'staff' or profile.status != 'pending':
            messages.error(request, '只能审核待审核状态的干事账号')
            return redirect('clubs:admin_dashboard')
    except UserProfile.DoesNotExist:
        messages.error(request, '用户角色信息不存在')
        return redirect('clubs:admin_dashboard')

    if request.method == 'POST':
        decision = request.POST.get('decision', '')
        review_comment = request.POST.get('review_comment', '').strip()

        if decision not in ['approved', 'rejected']:
            messages.error(request, '审核结果不合法')
            return redirect('clubs:review_staff_registration', user_id=user_id)

        # 更新用户状态
        profile.status = decision
        if review_comment:
            # 这里可以考虑将评论保存到其他字段或表中
            # 为了简单起见，我们暂时不保存评论
            pass
        profile.save()

        messages.success(request, f'用户 {user.username} 的注册申请已{'批准' if decision == 'approved' else '拒绝'}')
        return redirect('clubs:manage_users')

    context = {
        'user': user,
        'profile': profile,
    }
    return render(request, 'clubs/admin/review_staff_registration.html', context)



@login_required(login_url=settings.LOGIN_URL)
def admin_site_settings(request):
    """已合并到 manage_favicon，保留此视图以兼容旧链接"""
    return redirect('clubs:manage_favicon')


@login_required(login_url=settings.LOGIN_URL)
def manage_carousel(request):
    """管理轮播图列表"""
    if not _is_admin(request.user):
        messages.error(request, '仅管理员可以访问此页面')
        return redirect('clubs:index')

    carousel_images = CarouselImage.objects.all().order_by('-uploaded_at')
    return render(request, 'clubs/admin/manage_carousel.html', {
        'carousel_images': carousel_images,
    })


@login_required(login_url=settings.LOGIN_URL)
def add_carousel(request):
    """添加轮播图"""
    if not _is_admin(request.user):
        messages.error(request, '仅管理员可以访问此页面')
        return redirect('clubs:index')

    if request.method == 'POST':
        title = request.POST.get('title', '').strip()
        description = request.POST.get('description', '').strip()
        link = request.POST.get('link', '').strip()
        order = int(request.POST.get('order', 0))
        is_active = request.POST.get('is_active') == 'on'
        image = request.FILES.get('image')

        if not image:
            messages.error(request, '请选择要上传的图片')
            return render(request, 'clubs/admin/carousel_form.html')

        carousel = CarouselImage.objects.create(
            title=title,
            description=description,
            link=link,
            order=order,
            is_active=is_active,
            image=image,
            uploaded_by=request.user
        )
        messages.success(request, '轮播图添加成功')
        return redirect('clubs:manage_carousel')

    return render(request, 'clubs/admin/carousel_form.html')


@login_required(login_url=settings.LOGIN_URL)
def edit_carousel(request, carousel_id):
    """编辑轮播图"""
    if not _is_admin(request.user):
        messages.error(request, '仅管理员可以访问此页面')
        return redirect('clubs:index')

    carousel = get_object_or_404(CarouselImage, id=carousel_id)

    if request.method == 'POST':
        carousel.title = request.POST.get('title', '').strip()
        carousel.description = request.POST.get('description', '').strip()
        carousel.link = request.POST.get('link', '').strip()
        carousel.order = int(request.POST.get('order', 0))
        carousel.is_active = request.POST.get('is_active') == 'on'

        # 如果上传了新图片，替换旧图片
        new_image = request.FILES.get('image')
        if new_image:
            # 删除旧图片文件
            if carousel.image:
                try:
                    carousel.image.delete(save=False)
                except:
                    pass
            carousel.image = new_image

        carousel.save()
        messages.success(request, '轮播图更新成功')
        return redirect('clubs:manage_carousel')

    return render(request, 'clubs/admin/carousel_form.html', {
        'carousel': carousel,
    })


@login_required(login_url=settings.LOGIN_URL)
def delete_carousel(request, carousel_id):
    """删除轮播图"""
    if not _is_admin(request.user):
        messages.error(request, '仅管理员可以访问此页面')
        return redirect('clubs:index')

    carousel = get_object_or_404(CarouselImage, id=carousel_id)

    if request.method == 'POST':
        # 删除图片文件
        if carousel.image:
            try:
                carousel.image.delete(save=False)
            except:
                pass

        carousel.delete()
        messages.success(request, '轮播图删除成功')

    return redirect('clubs:manage_carousel')


@login_required(login_url=settings.LOGIN_URL)
def locked_accounts(request):
    """列出被锁定的用户名，供管理员解锁或重置密码"""
    if not _is_admin(request.user):
        messages.error(request, '仅管理员可访问此页面')
        return redirect('clubs:index')

    from django.core.cache import cache
    locked = []
    for u in User.objects.all():
        key = f'login_lock:user:{u.username}'
        if cache.get(key):
            locked.append(u)

    return render(request, 'clubs/admin/locked_accounts.html', {'locked': locked})

@login_required(login_url=settings.LOGIN_URL)
@require_http_methods(["GET", "POST"])
def publish_announcement(request):
    """发布公告 - 仅管理员可用"""
    if not _is_admin(request.user):
        messages.error(request, '仅管理员可以发布公告')
        return redirect('clubs:index')


# admin_force_reset_password 已移除
# 该功能由管理员界面的“重设密码”表单替代，移除以简化入口并减少重复功能。
# 如果将来需要恢复，再添加对应的视图和路由即可。


@login_required(login_url=settings.LOGIN_URL)
@require_http_methods(['POST'])
def unlock_account(request, username):
    """管理员解锁被锁账号（POST）"""
    if not _is_admin(request.user):
        messages.error(request, '仅管理员可执行此操作')
        return redirect('clubs:index')

    from django.core.cache import cache
    lock_key = f'login_lock:user:{username}'
    attempts_key = f'login_attempts:user:{username}'
    cache.delete(lock_key)
    cache.delete(attempts_key)

    messages.success(request, f'账号 {username} 已解锁')
    return redirect('clubs:locked_accounts')


@require_http_methods(["GET", "POST"])
def publish_announcement(request):
    """发布公告 - 仅管理员可用"""
    if not _is_admin(request.user):
        messages.error(request, '仅管理员可以发布公告')
        return redirect('clubs:index')

    if request.method == 'POST':
        title = request.POST.get('title', '').strip()
        content = request.POST.get('content', '').strip()
        status = request.POST.get('status', 'published')
        expires_at = request.POST.get('expires_at', '')
        attachment = request.FILES.get('attachment')

        errors = []
        if not title:
            errors.append('公告标题不能为空')
        if not content:
            errors.append('公告内容不能为空')

        if errors:
            announcements = Announcement.objects.all().order_by('-created_at')[:10]
            context = {
                'errors': errors,
                'title': title,
                'content': content,
                'announcements': announcements,
            }
            return render(request, 'clubs/admin/publish_announcement.html', context)

        announcement = Announcement.objects.create(
            title=title,
            content=content,
            status=status,
            created_by=request.user,
            published_at=timezone.now() if status == 'published' else None,
            expires_at=expires_at if expires_at else None,
            attachment=attachment,
        )

        messages.success(request, '公告发布成功！')
        return redirect('clubs:admin_dashboard')

    # GET 请求 - 获取最近的公告列表
    announcements = Announcement.objects.all().order_by('-created_at')[:10]
    context = {
        'announcements': announcements,
    }
    return render(request, 'clubs/admin/publish_announcement.html', context)


@login_required(login_url=settings.LOGIN_URL)
def delete_announcement(request, announcement_id):
    """删除公告 - 仅管理员可用"""
    if not _is_admin(request.user):
        messages.error(request, '仅管理员可以删除公告')
        return redirect('clubs:admin_dashboard')

    announcement = get_object_or_404(Announcement, pk=announcement_id)

    if request.method == 'POST':
        announcement_title = announcement.title
        announcement.delete()
        messages.success(request, f'公告"{announcement_title}"已删除')
        return redirect('clubs:admin_dashboard')

    # GET 请求：确认删除
    context = {
        'announcement': announcement,
    }
    return render(request, 'clubs/admin/confirm_delete_announcement.html', context)





@login_required(login_url=settings.LOGIN_URL)
@require_http_methods(["GET", "POST"])
def edit_announcement(request, announcement_id):
    """编辑公告 - 仅管理员可用"""
    if not _is_admin(request.user):
        messages.error(request, '仅管理员可以编辑公告')
        return redirect('clubs:admin_dashboard')

    announcement = get_object_or_404(Announcement, pk=announcement_id)

    if request.method == 'POST':
        title = request.POST.get('title', '').strip()
        content = request.POST.get('content', '').strip()
        status = request.POST.get('status', 'published')
        expires_at = request.POST.get('expires_at', '')
        attachment = request.FILES.get('attachment')

        errors = []
        if not title:
            errors.append('公告标题不能为空')
        if not content:
            errors.append('公告内容不能为空')

        if errors:
            context = {
                'errors': errors,
                'title': title,
                'content': content,
                'announcement': announcement,
            }
            return render(request, 'clubs/admin/edit_announcement.html', context)

        announcement.title = title
        announcement.content = content
        announcement.status = status
        announcement.expires_at = expires_at if expires_at else None

        # 如果有新附件，则更新附件
        if attachment:
            announcement.attachment = attachment

        # 如果状态从非发布状态变为发布状态，更新发布时间
        if status == 'published' and announcement.status != 'published':
            announcement.published_at = timezone.now()

        announcement.save()

        messages.success(request, '公告修改成功！')
        return redirect('clubs:admin_dashboard')

    # GET 请求 - 预填充表单
    context = {
        'announcement': announcement,
    }
    return render(request, 'clubs/admin/edit_announcement.html', context)


@login_required(login_url=settings.LOGIN_URL)
@require_http_methods(['GET'])
def download_user_import_template(request):
    """下载用户批量导入CSV模板（仅管理员/干事）。"""
    if not is_staff_or_admin(request.user):
        messages.error(request, '仅干事和管理员可以下载导入模板')
        return redirect('clubs:index')

    response = HttpResponse(content_type='text/csv; charset=utf-8-sig')
    response['Content-Disposition'] = 'attachment; filename="user_import_template.csv"'

    writer = csv.writer(response)
    writer.writerow(['用户名', '真实姓名', '邮箱', '电话', '微信', '学号', '角色', '密码', '部门', '政治面貌'])
    writer.writerow(['demo_president', '示例社长', 'demo@example.com', '13800138000', 'demo_wechat', '20260001', 'president', '123456', '', 'non_member'])
    writer.writerow(['demo_staff', '示例干事', 'staff@example.com', '13900139000', 'staff_wechat', '20260002', 'staff', '123456', '组织部', 'non_member'])

    return response


@login_required(login_url=settings.LOGIN_URL)
@require_http_methods(['GET'])
def export_all_users_and_clubs_csv(request):
    """导出全部用户与全部社团数据（ZIP内含两个CSV）。"""
    if not is_staff_or_admin(request.user):
        messages.error(request, '仅干事和管理员可以导出数据')
        return redirect('clubs:index')

    users_buffer = io.StringIO()
    users_writer = csv.writer(users_buffer)
    users_writer.writerow([
        '用户名', '真实姓名', '邮箱', '电话', '微信', '学号', '角色', '状态', '部门', '政治面貌', '加入时间'
    ])

    users_qs = User.objects.select_related('profile').all().order_by('id')
    for user in users_qs:
        profile = getattr(user, 'profile', None)
        users_writer.writerow([
            user.username,
            profile.get_full_name() if profile else user.first_name,
            user.email,
            profile.phone if profile else '',
            profile.wechat if profile else '',
            profile.student_id if profile else '',
            profile.role if profile else '',
            profile.status if profile else '',
            profile.department if profile else '',
            profile.political_status if profile else '',
            timezone.localtime(user.date_joined).strftime('%Y-%m-%d %H:%M:%S') if user.date_joined else '',
        ])

    clubs_buffer = io.StringIO()
    clubs_writer = csv.writer(clubs_buffer)
    clubs_writer.writerow([
        '社团ID', '社团名称', '状态', '成员数', '社长用户名', '社长姓名', '成立日期', '创建时间'
    ])

    clubs_qs = Club.objects.prefetch_related(
        Prefetch(
            'officers',
            queryset=Officer.objects.filter(position='president', is_current=True).select_related('user_profile__user'),
            to_attr='_president_list',
        )
    ).all().order_by('id')
    for club in clubs_qs:
        president_officer = club._president_list[0] if club._president_list else None
        president = president_officer.user_profile.user if president_officer and president_officer.user_profile else None
        president_name = ''
        if president_officer and president_officer.user_profile:
            president_name = president_officer.user_profile.get_full_name()
        clubs_writer.writerow([
            club.id,
            club.name,
            club.status,
            club.members_count,
            president.username if president else '',
            president_name,
            club.founded_date.strftime('%Y-%m-%d') if club.founded_date else '',
            timezone.localtime(club.created_at).strftime('%Y-%m-%d %H:%M:%S') if club.created_at else '',
        ])

    import zipfile

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        zip_file.writestr('all_users.csv', users_buffer.getvalue().encode('utf-8-sig'))
        zip_file.writestr('all_clubs.csv', clubs_buffer.getvalue().encode('utf-8-sig'))

    zip_buffer.seek(0)
    response = HttpResponse(zip_buffer.getvalue(), content_type='application/zip')
    response['Content-Disposition'] = 'attachment; filename="all_users_and_clubs_export.zip"'
    return response


def _csv_value(row, aliases):
    for key in aliases:
        value = row.get(key)
        if value is not None and str(value).strip() != '':
            return str(value).strip()
    return ''


@login_required(login_url=settings.LOGIN_URL)
@require_http_methods(['POST'])
def import_users_csv(request):
    """批量导入用户（仅CSV，仅管理员可用）。使用事务确保原子性，避免部分导入失败。"""
    is_ajax = request.headers.get('X-Requested-With') == 'XMLHttpRequest'

    def _json_error(msg, status=400):
        return JsonResponse({'success': False, 'message': msg}, status=status)

    if not _is_admin(request.user):
        if is_ajax:
            return _json_error('仅管理员可以批量导入用户', 403)
        messages.error(request, '仅管理员可以批量导入用户')
        return redirect('clubs:index')

    next_url = request.POST.get('next', '').strip() or request.META.get('HTTP_REFERER') or reverse('clubs:manage_users')
    uploaded = request.FILES.get('csv_file')
    if not uploaded:
        if is_ajax:
            return _json_error('请选择CSV文件后再导入')
        messages.error(request, '请选择CSV文件后再导入')
        return redirect(next_url)

    if not uploaded.name.lower().endswith('.csv'):
        if is_ajax:
            return _json_error('仅支持CSV文件导入')
        messages.error(request, '仅支持CSV文件导入')
        return redirect(next_url)

    raw_bytes = uploaded.read()
    text = None
    for encoding in ('utf-8-sig', 'utf-8', 'gbk'):
        try:
            text = raw_bytes.decode(encoding)
            break
        except Exception:
            continue

    if text is None:
        if is_ajax:
            return _json_error('CSV文件编码无法识别，请使用UTF-8编码')
        messages.error(request, 'CSV文件编码无法识别，请使用UTF-8编码')
        return redirect(next_url)

    reader = csv.DictReader(io.StringIO(text))
    if not reader.fieldnames:
        if is_ajax:
            return _json_error('CSV表头无效')
        messages.error(request, 'CSV表头无效')
        return redirect(next_url)

    created_users = 0
    updated_users = 0
    skipped = 0
    errors = []

    # 使用事务处理，确保导入失败时完全回滚，避免部分导入导致数据不一致
    try:
        from django.db import transaction
        with transaction.atomic():
            # 预处理：检查所有行的有效性和重复性，构建待创建/更新对象列表
            rows_to_process = []
            existing_student_ids = {}  # {student_id: username}

            # 加载数据库中已有的学号
            for profile in UserProfile.objects.filter(student_id__isnull=False).values('user__username', 'student_id'):
                if profile['student_id']:
                    existing_student_ids[profile['student_id']] = profile['user__username']

            # 跟踪本次导入中的学号，防止CSV内部重复
            import_student_ids = {}  # {student_id: row_idx}

            for idx, row in enumerate(reader, start=2):
                username = _csv_value(row, ['用户名', 'username'])
                real_name = _csv_value(row, ['真实姓名', 'real_name'])
                email = _csv_value(row, ['邮箱', 'email'])
                phone = _csv_value(row, ['电话', 'phone'])
                wechat = _csv_value(row, ['微信', 'wechat'])
                student_id = _csv_value(row, ['学号', 'student_id'])
                role = _csv_value(row, ['角色', 'role']).lower() or 'president'
                password = _csv_value(row, ['密码', 'password']) or '123456'
                department_name = _csv_value(row, ['部门', 'department'])
                political_status = _csv_value(row, ['政治面貌', 'political_status']) or 'non_member'

                if not username or not real_name:
                    skipped += 1
                    errors.append(f'第{idx}行缺少必填项（用户名/真实姓名）')
                    continue

                if role not in ['president', 'staff', 'admin']:
                    skipped += 1
                    errors.append(f'第{idx}行角色无效：{role}')
                    continue

                # 检查学号是否已存在于数据库（且不是同一个用户）
                if student_id:
                    if student_id in existing_student_ids and existing_student_ids[student_id] != username:
                        skipped += 1
                        errors.append(f'第{idx}行学号重复：{student_id}')
                        continue

                    # 检查本次导入中是否有重复学号
                    if student_id in import_student_ids:
                        skipped += 1
                        errors.append(f'第{idx}行学号与第{import_student_ids[student_id]}行重复')
                        continue

                    import_student_ids[student_id] = idx

                department_obj = None
                if department_name:
                    department_obj = Department.objects.filter(name=department_name).first()

                rows_to_process.append({
                    'idx': idx,
                    'username': username,
                    'real_name': real_name,
                    'email': email,
                    'phone': phone,
                    'wechat': wechat,
                    'student_id': student_id,
                    'role': role,
                    'password': password,
                    'department_name': department_name,
                    'department_obj': department_obj,
                    'political_status': political_status,
                })

            # 批量处理：查询所有现有用户
            existing_usernames = set(User.objects.filter(username__in=[r['username'] for r in rows_to_process]).values_list('username', flat=True))

            users_to_create = []
            users_to_update = []
            profiles_to_create = []
            profiles_to_update = []

            for row_data in rows_to_process:
                username = row_data['username']
                exists = username in existing_usernames

                if exists:
                    updated_users += 1
                    user = User.objects.get(username=username)
                    user.first_name = row_data['real_name']
                    if row_data['email']:
                        user.email = row_data['email']
                    user.set_password(row_data['password'])
                    users_to_update.append(user)
                else:
                    created_users += 1
                    user = User(
                        username=username,
                        first_name=row_data['real_name'],
                        email=row_data['email'] or ''
                    )
                    user.set_password(row_data['password'])
                    users_to_create.append(user)

            # 批量创建新用户
            if users_to_create:
                User.objects.bulk_create(users_to_create, ignore_conflicts=False)

            # 批量更新现有用户
            if users_to_update:
                User.objects.bulk_update(users_to_update, ['first_name', 'email', 'password'], batch_size=100)

            # 重新查询所有用户（因为新创建的用户需要id）
            user_map = {}
            for user in User.objects.filter(username__in=[r['username'] for r in rows_to_process]):
                user_map[user.username] = user

            # 构建Profile对象
            existing_profiles = {up.user_id: up for up in UserProfile.objects.filter(user_id__in=user_map.values())}

            for row_data in rows_to_process:
                user = user_map[row_data['username']]
                status = 'approved'  # 批量导入用户默认直接生效

                if user.id in existing_profiles:
                    profile = existing_profiles[user.id]
                    profile.role = row_data['role']
                    profile.status = status
                    profile.real_name = row_data['real_name']
                    profile.phone = row_data['phone']
                    profile.wechat = row_data['wechat']
                    profile.political_status = row_data['political_status']
                    profile.department = row_data['department_name'] if row_data['department_name'] else None
                    profile.department_link = row_data['department_obj']
                    if row_data['student_id']:
                        profile.student_id = row_data['student_id']
                    profile.must_change_password = row_data['password'] == '123456'
                    profiles_to_update.append(profile)
                else:
                    profile = UserProfile(
                        user=user,
                        role=row_data['role'],
                        status=status,
                        real_name=row_data['real_name'],
                        student_id=row_data['student_id'],
                        phone=row_data['phone'],
                        wechat=row_data['wechat'],
                        political_status=row_data['political_status'],
                        department=row_data['department_name'] if row_data['department_name'] else None,
                        department_link=row_data['department_obj'],
                        must_change_password=row_data['password'] == '123456',
                    )
                    profiles_to_create.append(profile)

            # 批量创建和更新Profile
            if profiles_to_create:
                UserProfile.objects.bulk_create(profiles_to_create, ignore_conflicts=False)

            if profiles_to_update:
                UserProfile.objects.bulk_update(
                    profiles_to_update,
                    ['role', 'status', 'real_name', 'phone', 'wechat', 'political_status', 'department', 'department_link', 'student_id', 'must_change_password'],
                    batch_size=100
                )

    except Exception as e:
        # 事务回滚，返回错误信息
        if is_ajax:
            return _json_error(f'导入失败：{str(e)}')
        messages.error(request, f'导入失败：{str(e)}')
        return redirect(next_url)

    summary_text = f'导入完成：新建{created_users}，更新{updated_users}，跳过{skipped}'
    if is_ajax:
        return JsonResponse({
            'success': True,
            'message': summary_text,
            'created_users': created_users,
            'updated_users': updated_users,
            'skipped': skipped,
            'errors': errors[:10],
            'next_url': next_url,
        })

    messages.success(request, summary_text)
    if errors:
        messages.warning(request, '部分数据有问题：' + '；'.join(errors[:10]))

    return redirect(next_url)


@login_required(login_url=settings.LOGIN_URL)
@require_http_methods(['GET'])
def download_club_import_template(request):
    """下载社团批量导入CSV模板（仅干事/管理员）。"""
    if not is_staff_or_admin(request.user):
        messages.error(request, '仅干事和管理员可以下载社团导入模板')
        return redirect('clubs:index')

    response = HttpResponse(content_type='text/csv; charset=utf-8-sig')
    response['Content-Disposition'] = 'attachment; filename="club_import_template.csv"'

    writer = csv.writer(response)
    writer.writerow(['社团名称', '社团简介', '成立日期', '状态', '成员数', '社长用户名'])
    writer.writerow(['示例社团A', '示例社团简介', '2024-09-01', 'active', '36', 'demo_president'])
    writer.writerow(['示例社团B', '可为空', '', 'inactive', '0', ''])
    return response


@login_required(login_url=settings.LOGIN_URL)
@require_http_methods(['POST'])
def import_clubs_csv(request):
    """批量导入社团（仅CSV，干事/管理员可用）。"""
    if not is_staff_or_admin(request.user):
        messages.error(request, '仅干事和管理员可以批量导入社团')
        return redirect('clubs:index')

    next_url = request.POST.get('next', '').strip() or request.META.get('HTTP_REFERER') or reverse('clubs:staff_management')
    uploaded = request.FILES.get('csv_file')
    if not uploaded:
        messages.error(request, '请选择CSV文件后再导入')
        return redirect(next_url)

    if not uploaded.name.lower().endswith('.csv'):
        messages.error(request, '仅支持CSV文件导入')
        return redirect(next_url)

    raw_bytes = uploaded.read()
    text = None
    for encoding in ('utf-8-sig', 'utf-8', 'gbk'):
        try:
            text = raw_bytes.decode(encoding)
            break
        except Exception:
            continue

    if text is None:
        messages.error(request, 'CSV文件编码无法识别，请使用UTF-8编码')
        return redirect(next_url)

    reader = csv.DictReader(io.StringIO(text))
    if not reader.fieldnames:
        messages.error(request, 'CSV表头无效')
        return redirect(next_url)

    created_clubs = 0
    updated_clubs = 0
    skipped = 0
    errors = []

    valid_status = {'active', 'inactive', 'suspended'}

    for idx, row in enumerate(reader, start=2):
        name = _csv_value(row, ['社团名称', 'name'])
        description = _csv_value(row, ['社团简介', 'description'])
        founded_date_raw = _csv_value(row, ['成立日期', 'founded_date'])
        status = (_csv_value(row, ['状态', 'status']).lower() or 'active')
        members_raw = _csv_value(row, ['成员数', 'members_count'])
        president_username = _csv_value(row, ['社长用户名', 'president_username'])

        if not name:
            skipped += 1
            errors.append(f'第{idx}行缺少必填项（社团名称）')
            continue

        if status not in valid_status:
            skipped += 1
            errors.append(f'第{idx}行状态无效：{status}')
            continue

        founded_date = None
        if founded_date_raw:
            try:
                founded_date = datetime.strptime(founded_date_raw, '%Y-%m-%d').date()
            except ValueError:
                skipped += 1
                errors.append(f'第{idx}行成立日期格式错误，应为YYYY-MM-DD')
                continue
        else:
            founded_date = timezone.now().date()

        members_count = 0
        if members_raw:
            try:
                members_count = int(members_raw)
                if members_count < 0:
                    raise ValueError('negative')
            except Exception:
                skipped += 1
                errors.append(f'第{idx}行成员数无效：{members_raw}')
                continue

        president = None
        if president_username:
            president = User.objects.filter(username=president_username).first()
            if president is None:
                skipped += 1
                errors.append(f'第{idx}行社长用户名不存在：{president_username}')
                continue

        club, created = Club.objects.get_or_create(
            name=name,
            defaults={
                'description': description,
                'founded_date': founded_date,
                'status': status,
                'members_count': members_count,
            },
        )

        if created:
            created_clubs += 1
        else:
            updated_clubs += 1
            club.description = description
            club.founded_date = founded_date
            club.status = status
            club.members_count = members_count
            club.save()

        # 设置/更新社长 Officer 记录
        if president:
            try:
                president_profile = president.profile
                Officer.objects.filter(
                    club=club, position='president', is_current=True
                ).exclude(user_profile=president_profile).update(
                    is_current=False, end_date=timezone.now().date()
                )
                Officer.objects.update_or_create(
                    club=club,
                    user_profile=president_profile,
                    position='president',
                    defaults={'is_current': True, 'appointed_date': timezone.now().date(), 'end_date': None}
                )
            except Exception:
                errors.append(f'第{idx}行社长 Officer 记录更新失败')

    messages.success(request, f'社团导入完成：新建{created_clubs}，更新{updated_clubs}，跳过{skipped}')
    if errors:
        messages.warning(request, '部分数据有问题：' + '；'.join(errors[:5]))

    return redirect(next_url)


@login_required(login_url=settings.LOGIN_URL)
def manage_users(request):
    """用户管理 - 仅管理员可用"""
    if not _is_admin(request.user):
        messages.error(request, '仅管理员可以管理用户')
        return redirect('clubs:index')

    if request.method == 'POST' and request.POST.get('action') == 'delete_user':
        user_id = request.POST.get('user_id', '').strip()
        target_user = get_object_or_404(User, pk=user_id)

        if target_user == request.user:
            messages.error(request, '不能删除当前登录管理员账号')
        elif target_user.is_superuser:
            messages.error(request, '不能删除超级管理员账号')
        else:
            username = target_user.username
            target_user.delete()
            messages.success(request, f'已删除用户账号：{username}')

        return redirect('clubs:manage_users')

    # 获取所有用户，使用select_related加载关联的UserProfile以包含状态信息
    # 便于管理员审核待审核的干事账号
    users = User.objects.select_related('profile').all()

    # 搜索过滤
    search = request.GET.get('search', '').strip()
    if search:
        from django.db.models import Q
        users = users.filter(Q(username__icontains=search) | Q(email__icontains=search))

    # 角色过滤
    role = request.GET.get('role', '').strip()
    if role:
        users = users.filter(profile__role=role)

    # 检查哪些用户被锁（用于在用户列表中显示解锁按钮）
    from django.core.cache import cache
    locked_usernames = set()
    for u in User.objects.all():
        key = f'login_lock:user:{u.username}'
        if cache.get(key):
            locked_usernames.add(u.username)

    context = {
        'users': users,
        'total_users': User.objects.count(),
        'search': search,
        'role': role,
        'locked_usernames': locked_usernames,
    }
    return render(request, 'clubs/admin/manage_users.html', context)


@login_required(login_url=settings.LOGIN_URL)
def staff_view_users(request):
    """干事查看用户列表 - 查看专用，无编辑权限"""
    if not is_staff_or_admin(request.user):
        messages.error(request, '仅干事可以查看用户列表')
        return redirect('clubs:index')

    # 获取所有用户，但不提供编辑功能
    users = User.objects.select_related('profile').all()

    # 搜索过滤
    search = request.GET.get('search', '').strip()
    if search:
        from django.db.models import Q
        users = users.filter(Q(username__icontains=search) | Q(email__icontains=search))

    # 角色过滤
    role = request.GET.get('role', '').strip()
    if role:
        users = users.filter(profile__role=role)

    context = {
        'users': users,
        'total_users': User.objects.count(),
        'search': search,
        'role': role,
        'is_staff_view': True,  # 标记为干事视图
    }
    return render(request, 'clubs/staff/view_users.html', context)


@login_required(login_url=settings.LOGIN_URL)
@require_http_methods(["GET", "POST"])
def create_user(request):
    """管理员创建用户账户 - 仅管理员可用"""
    if not _is_admin(request.user):
        messages.error(request, '仅管理员可以创建用户账户')
        return redirect('clubs:index')

    if request.method == 'POST':
        username = request.POST.get('username', '').strip()
        real_name = request.POST.get('real_name', '').strip()
        email = request.POST.get('email', '').strip()
        phone = request.POST.get('phone', '').strip()
        wechat = request.POST.get('wechat', '').strip()
        password = request.POST.get('password', '').strip()
        password2 = request.POST.get('password2', '').strip()
        role = request.POST.get('role', 'president').strip()
        student_id = request.POST.get('student_id', '').strip()

        errors = []

        # 验证
        if not username:
            errors.append('登录用户名不能为空')
        elif User.objects.filter(username=username).exists():
            errors.append('登录用户名已存在')
        elif len(username) < 3 or len(username) > 30:
            errors.append('登录用户名长度应在3-30个字符之间')

        if not real_name:
            errors.append('姓名不能为空')

        if not email:
            errors.append('邮箱不能为空')
        elif User.objects.filter(email=email).exists():
            errors.append('邮箱已被使用')

        if not phone:
            errors.append('电话不能为空')

        if not wechat:
            errors.append('微信号不能为空')

        if not password:
            errors.append('密码不能为空')
        elif len(password) < 6:
            errors.append('密码至少6个字符')

        if password != password2:
            errors.append('两次输入的密码不一致')

        if role not in ['president', 'staff', 'admin']:
            errors.append('无效的用户角色')

        # 学号验证（必填字段）
        if not student_id:
            errors.append('学号不能为空')

        if errors:
            context = {
                'errors': errors,
                'form_data': request.POST,
            }
            return render(request, 'clubs/admin/create_user.html', context)

        # 创建用户
        user = User.objects.create_user(
            username=username,
            email=email,
            password=password,
            first_name=real_name
        )

        # 确定用户状态 - 干事默认待审核，其他类型直接批准
        status = 'pending' if role == 'staff' else 'approved'

        # 创建用户角色信息
        profile, created = UserProfile.objects.get_or_create(
            user=user,
            defaults={
                'role': role,
                'status': status,
                'real_name': real_name,
                'phone': phone,
                'wechat': wechat,
                'student_id': student_id
            }
        )

        # 如果profile已经存在，更新它
        if not created:
            profile.role = role
            profile.status = status
            profile.real_name = real_name
            profile.phone = phone
            profile.wechat = wechat
            profile.student_id = student_id
            profile.save()

        role_display = dict(UserProfile.ROLE_CHOICES).get(role, role)
        messages.success(request, f'成功创建用户账户：{username}（角色：{role_display}）')
        return redirect('clubs:manage_users')

    return render(request, 'clubs/admin/create_user.html')



@login_required(login_url=settings.LOGIN_URL)
@require_http_methods(["GET", "POST"])
def admin_edit_user_account(request, user_id):
    """管理员编辑用户账户信息 - 仅管理员可用"""
    if not _is_admin(request.user):
        messages.error(request, '仅管理员可以编辑用户账户信息')
        return redirect('clubs:index')

    # 获取目标用户
    target_user = get_object_or_404(User, pk=user_id)

    # 确保管理员不能编辑自己的账户（使用自己的账户设置页面）
    if request.user == target_user:
        messages.error(request, '请使用您自己的账户设置页面编辑个人信息')
        return redirect('clubs:change_account_settings')

    errors = []
    success_messages = []

    if request.method == 'POST':
        action = request.POST.get('action', '')

        # 修改用户名
        if action == 'change_username':
            new_username = request.POST.get('new_username', '').strip()

            if not new_username:
                errors.append('新用户名不能为空')
            elif len(new_username) < 3:
                errors.append('用户名至少3个字符')
            elif User.objects.exclude(id=target_user.id).filter(username=new_username).exists():
                errors.append('用户名已被使用')
            else:
                old_username = target_user.username
                target_user.username = new_username
                target_user.save()
                success_messages.append(f'已将用户 {old_username} 的用户名修改为 {new_username}')

        # 修改密码
        elif action == 'change_password':
            new_password = request.POST.get('new_password', '').strip()
            confirm_password = request.POST.get('confirm_password', '').strip()

            if not new_password:
                errors.append('新密码不能为空')
            elif len(new_password) < 6:
                errors.append('新密码至少6个字符')
            elif new_password != confirm_password:
                errors.append('两次密码不一致')
            else:
                target_user.set_password(new_password)
                target_user.save()
                success_messages.append(f'已成功重置用户 {target_user.username} 的密码')

        # 修改角色
        elif action == 'change_role':
            new_role = request.POST.get('new_role', '')

            if new_role not in ['president', 'staff', 'admin', 'member']:
                errors.append('角色不合法')
            else:
                try:
                    profile = target_user.profile
                    old_role = profile.get_role_display()
                    profile.role = new_role
                    profile.save()
                    role_display = dict(UserProfile.ROLE_CHOICES).get(new_role, new_role)
                    success_messages.append(f'已将用户 {target_user.username} 的角色从「{old_role}」更改为「{role_display}」')
                except UserProfile.DoesNotExist:
                    # 创建UserProfile时处理student_id唯一性约束
                    import uuid
                    unique_student_id = f"USER_{target_user.username}_{str(uuid.uuid4())[:8]}"

                    profile = UserProfile.objects.create(
                        user=target_user,
                        role=new_role,
                        student_id=unique_student_id  # 设置唯一的student_id
                    )
                    role_display = dict(UserProfile.ROLE_CHOICES).get(new_role, new_role)
                    success_messages.append(f'已为用户 {target_user.username} 创建角色：「{role_display}」')

        # 修改全名
        elif action == 'change_full_name':
            full_name = request.POST.get('full_name', '').strip()

            try:
                profile = target_user.profile
                profile.real_name = full_name
                profile.save()
                success_messages.append(f'已更新用户 {target_user.username} 的全名：{full_name or "（已清空）"}')
            except UserProfile.DoesNotExist:
                import uuid
                unique_student_id = f"USER_{target_user.username}_{str(uuid.uuid4())[:8]}"
                profile = UserProfile.objects.create(
                    user=target_user,
                    role='member',
                    real_name=full_name,
                    student_id=unique_student_id
                )
                success_messages.append(f'已为用户 {target_user.username} 设置全名：{full_name}')

        # 修改邮箱
        elif action == 'change_email':
            email = request.POST.get('email', '').strip()

            if not email:
                errors.append('邮箱不能为空')
            elif not email.count('@'):
                errors.append('请输入有效的邮箱地址')
            elif User.objects.exclude(id=target_user.id).filter(email=email).exists():
                errors.append('此邮箱已被其他用户使用')
            else:
                target_user.email = email
                target_user.save()
                success_messages.append(f'已更新用户 {target_user.username} 的邮箱：{email}')

        # 修改电话
        elif action == 'change_phone':
            phone = request.POST.get('phone', '').strip()

            try:
                profile = target_user.profile
                profile.phone = phone
                profile.save()
                success_messages.append(f'已更新用户 {target_user.username} 的电话：{phone or "（已清空）"}')
            except UserProfile.DoesNotExist:
                import uuid
                unique_student_id = f"USER_{target_user.username}_{str(uuid.uuid4())[:8]}"
                profile = UserProfile.objects.create(
                    user=target_user,
                    role='member',
                    phone=phone,
                    student_id=unique_student_id
                )
                success_messages.append(f'已为用户 {target_user.username} 设置电话：{phone}')

        # 修改学号
        elif action == 'change_student_id':
            student_id = request.POST.get('student_id', '').strip()

            try:
                profile = target_user.profile
                profile.student_id = student_id
                profile.save()
                success_messages.append(f'已更新用户 {target_user.username} 的学号：{student_id or "（已清空）"}')
            except UserProfile.DoesNotExist:
                import uuid
                unique_student_id = student_id or f"USER_{target_user.username}_{str(uuid.uuid4())[:8]}"
                profile = UserProfile.objects.create(
                    user=target_user,
                    role='member',
                    student_id=unique_student_id
                )
                success_messages.append(f'已为用户 {target_user.username} 设置学号：{student_id}')

        # 修改QQ
        elif action == 'change_qq':
            qq = request.POST.get('qq', '').strip()

            try:
                profile = target_user.profile
                profile.qq = qq
                profile.save()
                success_messages.append(f'已更新用户 {target_user.username} 的QQ：{qq or "（已清空）"}')
            except UserProfile.DoesNotExist:
                import uuid
                unique_student_id = f"USER_{target_user.username}_{str(uuid.uuid4())[:8]}"
                profile = UserProfile.objects.create(
                    user=target_user,
                    role='member',
                    qq=qq,
                    student_id=unique_student_id
                )
                success_messages.append(f'已为用户 {target_user.username} 设置QQ：{qq}')

        # 修改微信
        elif action == 'change_wechat':
            wechat = request.POST.get('wechat', '').strip()

            try:
                profile = target_user.profile
                profile.wechat = wechat
                profile.save()
                success_messages.append(f'已更新用户 {target_user.username} 的微信：{wechat or "（已清空）"}')
            except UserProfile.DoesNotExist:
                import uuid
                unique_student_id = f"USER_{target_user.username}_{str(uuid.uuid4())[:8]}"
                profile = UserProfile.objects.create(
                    user=target_user,
                    role='member',
                    wechat=wechat,
                    student_id=unique_student_id
                )
                success_messages.append(f'已为用户 {target_user.username} 设置微信：{wechat}')

        # 修改性别
        elif action == 'change_gender':
            gender = request.POST.get('gender', '').strip()

            if gender and gender not in ['male', 'female', 'other']:
                errors.append('性别选项不合法')
            else:
                try:
                    profile = target_user.profile
                    profile.gender = gender
                    profile.save()
                    gender_display = dict(UserProfile.GENDER_CHOICES).get(gender, '不设定') if gender else '不设定'
                    success_messages.append(f'已更新用户 {target_user.username} 的性别：{gender_display}')
                except UserProfile.DoesNotExist:
                    import uuid
                    unique_student_id = f"USER_{target_user.username}_{str(uuid.uuid4())[:8]}"
                    profile = UserProfile.objects.create(
                        user=target_user,
                        role='member',
                        gender=gender,
                        student_id=unique_student_id
                    )
                    gender_display = dict(UserProfile.GENDER_CHOICES).get(gender, '不设定') if gender else '不设定'
                    success_messages.append(f'已为用户 {target_user.username} 设置性别：{gender_display}')

        # 修改用户详细信息
        elif action == 'change_user_info':
            real_name = request.POST.get('real_name', '').strip()
            student_id = request.POST.get('student_id', '').strip()
            phone = request.POST.get('phone', '').strip()
            wechat = request.POST.get('wechat', '').strip()
            political_status = request.POST.get('political_status', 'non_member')
            email = request.POST.get('email', '').strip()

            # 验证必填字段
            if not email:
                errors.append('邮箱不能为空')
            elif User.objects.exclude(id=target_user.id).filter(email=email).exists():
                errors.append('邮箱已被使用')

            if not real_name:
                errors.append('真实姓名不能为空')

            if student_id and UserProfile.objects.exclude(user=target_user).filter(student_id=student_id).exists():
                errors.append('学号已被使用')

            if not errors:
                # 更新用户基本信息
                target_user.email = email
                target_user.save()

                # 更新用户资料信息
                try:
                    profile = target_user.profile
                    profile.real_name = real_name
                    profile.student_id = student_id
                    profile.phone = phone
                    profile.wechat = wechat
                    profile.political_status = political_status
                    profile.save()
                except UserProfile.DoesNotExist:
                    # 如果用户没有资料，创建一个
                    import uuid
                    unique_student_id = student_id or f"USER_{target_user.username}_{str(uuid.uuid4())[:8]}"

                    profile = UserProfile.objects.create(
                        user=target_user,
                        role='president',  # 默认角色
                        real_name=real_name,
                        student_id=unique_student_id,
                        phone=phone,
                        wechat=wechat,
                        political_status=political_status
                    )

                success_messages.append(f'已成功更新用户 {target_user.username} 的详细信息')

        # 删除用户
        elif action == 'delete_user':
            if target_user.is_superuser:
                errors.append('不能删除超级管理员账户')
            else:
                username = target_user.username
                try:
                    target_user.delete()
                    messages.success(request, f'已删除用户账户：{username}')
                    return redirect('clubs:manage_users')
                except Exception as e:
                    errors.append(f'删除失败：{str(e)}')

        # 切换账户启用状态
        elif action == 'toggle_active':
            if target_user.is_superuser:
                errors.append('不能禁用超级管理员账户')
            else:
                old_state = '启用' if target_user.is_active else '禁用'
                target_user.is_active = not target_user.is_active
                target_user.save(update_fields=['is_active'])

                # 同步生命周期状态：禁用即进入不活跃计时，启用则恢复活跃。
                try:
                    profile = target_user.profile
                    if target_user.is_active:
                        profile.account_status = 'active'
                        profile.inactive_since = None
                        if profile.status == 'inactive':
                            profile.status = 'approved'
                        profile.save(update_fields=['account_status', 'inactive_since', 'status', 'updated_at'])
                    else:
                        mark_profile_inactive(profile, reason='admin_disable')
                except UserProfile.DoesNotExist:
                    pass

                new_state = '启用' if target_user.is_active else '禁用'
                success_messages.append(f'已将用户 {target_user.username} 的账户状态从「{old_state}」更改为「{new_state}」')


    # 获取用户角色信息
    try:
        profile = target_user.profile
    except UserProfile.DoesNotExist:
        profile = None

    context = {
        'target_user': target_user,
        'profile': profile,
        'errors': errors,
        'success_messages': success_messages,
        'is_admin_view': True,
        'ROLE_CHOICES': UserProfile.ROLE_CHOICES,
        'POLITICAL_STATUS_CHOICES': UserProfile.POLITICAL_STATUS_CHOICES,
    }

    return render(request, 'clubs/admin/admin_edit_user_account.html', context)


@login_required(login_url=settings.LOGIN_URL)
@require_http_methods(["GET", "POST"])
def change_user_role(request, user_id):
    """修改用户角色 - 仅管理员可用"""
    if not _is_admin(request.user):
        messages.error(request, '仅管理员可以修改用户角色')
        return redirect('clubs:index')

    target_user = get_object_or_404(User, pk=user_id)

    if request.method == 'POST':
        new_role = request.POST.get('new_role', '')

        if new_role not in ['president', 'staff', 'admin']:
            messages.error(request, '角色不合法')
            context = {'user': target_user}
            return render(request, 'clubs/admin/change_user_role.html', context)

        # 干事注册需要管理员同意（已确保只有管理员可以执行此操作）

        try:
            profile = target_user.profile
            old_role = profile.get_role_display()
            profile.role = new_role
            profile.save()
            role_display = dict(UserProfile.ROLE_CHOICES).get(new_role, new_role)
            messages.success(request, f'已将 {target_user.username} 的角色从「{old_role}」更改为「{role_display}」')
        except UserProfile.DoesNotExist:
            # 创建UserProfile时处理student_id唯一性约束
            # 为新创建的用户生成一个唯一的student_id，使用用户名+时间戳确保唯一性
            import uuid
            unique_student_id = f"USER_{target_user.username}_{str(uuid.uuid4())[:8]}"

            profile = UserProfile.objects.create(
                user=target_user,
                role=new_role,
                student_id=unique_student_id  # 设置唯一的student_id
            )
            role_display = dict(UserProfile.ROLE_CHOICES).get(new_role, new_role)
            messages.success(request, f'已为 {target_user.username} 创建角色：「{role_display}」')

        return redirect('clubs:manage_users')

    # GET 请求
    try:
        profile = target_user.profile
    except UserProfile.DoesNotExist:
        profile = None

    context = {
        'user': target_user,
        'profile': profile,
    }
    return render(request, 'clubs/admin/change_user_role.html', context)


@login_required(login_url=settings.LOGIN_URL)
@require_http_methods(['GET', 'POST'])
def change_staff_attributes(request, user_id):
    """修改干事的部门和职级属性 - 仅管理员可用"""
    if not _is_admin(request.user):
        messages.error(request, '仅管理员可以修改干事属性')
        return redirect('clubs:index')

    target_user = get_object_or_404(User, pk=user_id)

    # 检查目标用户是否为干事或管理员
    try:
        profile = target_user.profile
        if profile.role not in ('staff', 'admin'):
            messages.error(request, '仅干事或管理员支持修改部门/职级属性')
            return redirect('clubs:manage_users')
    except UserProfile.DoesNotExist:
        messages.error(request, '用户角色信息不存在')
        return redirect('clubs:manage_users')

    if request.method == 'POST':
        department = request.POST.get('department', '').strip()
        staff_level = request.POST.get('staff_level', '').strip()

        # 验证部门和职级
        valid_levels = dict(UserProfile.STAFF_LEVEL_CHOICES).keys()

        selected_department = None
        if department:
            try:
                selected_department = Department.objects.get(id=int(department))
            except (ValueError, Department.DoesNotExist):
                messages.error(request, '部门选择无效')
                return redirect('clubs:manage_users')

        if staff_level and staff_level not in valid_levels:
            messages.error(request, '职级选择无效')
            return redirect('clubs:manage_users')

        try:
            old_department = profile.department_link.name if profile.department_link else (profile.department or '未设定')
            old_level = profile.get_staff_level_display() if profile.staff_level else '未设定'

            profile.department_link = selected_department
            profile.department = selected_department.name if selected_department else None
            profile.staff_level = staff_level if staff_level else profile.staff_level
            profile.save()

            new_department = profile.department_link.name if profile.department_link else (profile.department or '未设定')
            new_level = profile.get_staff_level_display()

            messages.success(request, f'已修改 {target_user.username} 的干事属性：部门「{old_department}」→「{new_department}」，职级「{old_level}」→「{new_level}」')
        except Exception as e:
            messages.error(request, f'修改失败：{str(e)}')

        return redirect('clubs:manage_users')

    context = {
        'user': target_user,
        'profile': profile,
        'departments': Department.objects.all().order_by('order', 'name'),
        'staff_level_choices': UserProfile.STAFF_LEVEL_CHOICES,
    }
    return render(request, 'clubs/admin/change_staff_attributes.html', context)




@login_required(login_url=settings.LOGIN_URL)
@require_http_methods(['GET', 'POST'])
def manage_smtp_config(request):
    if not _is_admin(request.user):
        messages.error(request, '只有管理员可以管理 SMTP 配置')
        return redirect('clubs:index')

    if request.method == 'POST':
        action = request.POST.get('action', '')
        config_id = request.POST.get('config_id', '')
        if action in ['create', 'edit']:
            config = get_object_or_404(SMTPConfig, pk=config_id) if action == 'edit' else SMTPConfig()
            provider = request.POST.get('provider', '').strip()
            smtp_host = request.POST.get('smtp_host', '').strip()
            smtp_port = request.POST.get('smtp_port', '').strip()
            sender_email = request.POST.get('sender_email', '').strip()
            sender_password = request.POST.get('sender_password', '').strip()
            if not provider or not smtp_host or not smtp_port or not sender_email:
                messages.error(request, '请完整填写 SMTP 配置')
                return redirect('clubs:manage_smtp_config')
            try:
                config.smtp_port = int(smtp_port)
            except ValueError:
                messages.error(request, 'SMTP 端口必须是数字')
                return redirect('clubs:manage_smtp_config')
            if request.POST.get('is_active') == 'on':
                SMTPConfig.objects.all().update(is_active=False)
            config.provider = provider
            config.smtp_host = smtp_host
            config.sender_email = sender_email
            if sender_password:
                config.sender_password = sender_password
            config.use_tls = request.POST.get('use_tls') == 'on'
            config.is_active = request.POST.get('is_active') == 'on'
            config.save()
            messages.success(request, 'SMTP 配置已保存')
            return redirect('clubs:manage_smtp_config')
        if action == 'delete':
            get_object_or_404(SMTPConfig, pk=config_id).delete()
            messages.success(request, 'SMTP 配置已删除')
            return redirect('clubs:manage_smtp_config')
        if action == 'activate':
            SMTPConfig.objects.all().update(is_active=False)
            config = get_object_or_404(SMTPConfig, pk=config_id)
            config.is_active = True
            config.save(update_fields=['is_active'])
            messages.success(request, 'SMTP 配置已启用')
            return redirect('clubs:manage_smtp_config')
        if action == 'test_email':
            config = get_object_or_404(SMTPConfig, pk=config_id)
            test_email = request.POST.get('test_email', '').strip()
            if not test_email:
                messages.error(request, '请填写测试收件邮箱')
                return redirect('clubs:manage_smtp_config')
            from .email_utils import send_test_email_with_config
            success, msg = send_test_email_with_config(config, test_email)
            messages.success(request, msg) if success else messages.error(request, msg)
            return redirect('clubs:manage_smtp_config')

    return render(request, 'clubs/admin/smtp_config.html', {'configs': SMTPConfig.objects.all()})




def manage_storage_config(request):
    """存储后端配置页（S3 / 本地切换）。

    支持操作：
      * save         保存当前配置
      * test         测试 S3 连接（不持久化）
      * reset_local  紧急回退到本地存储
    """
    if not _is_admin(request.user):
        messages.error(request, '只有管理员可以管理存储配置')
        return redirect('clubs:index')

    from .models import StorageConfig
    from .storage_backends import ClubStorage, cleanup_temp_files

    config = StorageConfig.get_active_config()
    test_result = None
    errors = []

    if request.method == 'POST':
        action = request.POST.get('action', 'save')

        if action == 'reset_local':
            config.backend_type = 'local'
            config.is_active = True
            config.save()
            messages.success(request, '已紧急回退到本地存储')
            return redirect('clubs:manage_storage_config')

        if action == 'test':
            # 仅测试，不保存
            test_cfg = {
                's3_endpoint_url': request.POST.get('s3_endpoint_url', '').strip(),
                's3_region': request.POST.get('s3_region', '').strip(),
                's3_bucket_name': request.POST.get('s3_bucket_name', '').strip(),
                's3_access_key_id': request.POST.get('s3_access_key_id', '').strip(),
                's3_secret_access_key': request.POST.get('s3_secret_access_key', '').strip(),
                's3_addressing_style': request.POST.get('s3_addressing_style', 'auto').strip(),
            }
            if not test_cfg['s3_bucket_name'] or not test_cfg['s3_access_key_id'] \
                    or not test_cfg['s3_secret_access_key']:
                errors.append('测试 S3 连接时必须填写 bucket / AK / SK')
                test_result = (False, '缺少必填字段')
            else:
                storage = ClubStorage()
                ok, msg = storage.test_s3_connection(test_cfg)
                test_result = (ok, msg)
                if ok:
                    messages.success(request, msg)
                else:
                    errors.append(msg)
                # 回显表单
                return render(request, 'clubs/admin/storage_config.html', {
                    'config': config,
                    'test_result': test_result,
                    'errors': errors,
                    'test_form': test_cfg,
                })

        if action == 'save':
            backend_type = request.POST.get('backend_type', 'local').strip()
            if backend_type not in ('local', 's3'):
                errors.append('后端类型必须是 local 或 s3')
            else:
                config.backend_type = backend_type
                config.is_active = request.POST.get('is_active') == 'on'
                config.s3_endpoint_url = request.POST.get('s3_endpoint_url', '').strip()
                config.s3_region = request.POST.get('s3_region', '').strip()
                config.s3_bucket_name = request.POST.get('s3_bucket_name', '').strip()
                config.s3_access_key_id = request.POST.get('s3_access_key_id', '').strip()
                # SK 留空表示保持原值
                new_sk = request.POST.get('s3_secret_access_key', '').strip()
                if new_sk:
                    config.s3_secret_access_key = new_sk
                elif not config.s3_secret_access_key and backend_type == 's3':
                    errors.append('切换到 S3 时必须填写 Secret Access Key')
                config.s3_custom_domain = request.POST.get('s3_custom_domain', '').strip()
                config.s3_addressing_style = request.POST.get('s3_addressing_style', 'auto').strip()
                config.s3_use_path_style = request.POST.get('s3_use_path_style') == 'on'
                try:
                    config.presigned_url_expiration = int(
                        request.POST.get('presigned_url_expiration', '3600') or 3600
                    )
                except ValueError:
                    errors.append('预签名有效期必须是数字')

                # 切换到 S3 前先测试连接，避免配置错误导致后续保存失败
                if backend_type == 's3' and not errors:
                    storage = ClubStorage()
                    ok, msg = storage.test_s3_connection({
                        's3_endpoint_url': config.s3_endpoint_url,
                        's3_region': config.s3_region,
                        's3_bucket_name': config.s3_bucket_name,
                        's3_access_key_id': config.s3_access_key_id,
                        's3_secret_access_key': config.s3_secret_access_key,
                        's3_addressing_style': config.s3_addressing_style,
                    })
                    if not ok:
                        errors.append('S3 连接测试失败：%s' % msg)
                    else:
                        messages.success(request, 'S3 连接测试通过')

                if not errors:
                    config.save()
                    messages.success(request, '存储配置已保存')
                    return redirect('clubs:manage_storage_config')

    return render(request, 'clubs/admin/storage_config.html', {
        'config': config,
        'test_result': test_result,
        'errors': errors,
        'test_form': None,
    })




def _default_cycle_name(channel):
    today = timezone.localdate()
    if channel.cycle_type == 'year':
        return f'{today.year}年'
    if channel.cycle_type == 'month':
        return f'{today.year}年{today.month:02d}月'
    if channel.cycle_type == 'day':
        return today.strftime('%Y-%m-%d')
    return ''


@login_required(login_url=settings.LOGIN_URL)
@require_http_methods(['POST'])
def toggle_form_channel_cycle(request, channel_id):
    if not is_staff_or_admin(request.user):
        return HttpResponseForbidden("No permission")

    channel = get_object_or_404(FormChannel, pk=channel_id, allow_staff_toggle=True)
    needs_cycle = channel.cycle_type != 'none' or channel.submission_policy == 'once_per_cycle'
    if needs_cycle:
        channel.is_active = True
        channel.submission_policy = 'once_per_cycle'
        channel.save(update_fields=['is_active', 'submission_policy', 'updated_at'])
        active_cycle = channel.cycles.filter(is_active=True).order_by('-sequence', '-starts_at').first()
        if active_cycle:
            active_cycle.is_active = False
            active_cycle.ends_at = timezone.now()
            active_cycle.save(update_fields=['is_active', 'ends_at'])
            messages.success(request, f"{channel.name} 已关闭：{active_cycle.name}")
        else:
            cycle = create_form_cycle(channel, name=_default_cycle_name(channel), user=request.user)
            messages.success(request, f"{channel.name} 已开启：{cycle.name}")
    else:
        channel.is_active = not channel.is_active
        channel.save(update_fields=['is_active', 'updated_at'])
        messages.success(request, f"{channel.name} 已{'开启' if channel.is_active else '关闭'}")
    return redirect(request.META.get('HTTP_REFERER') or reverse('clubs:staff_management'))


@login_required(login_url=settings.LOGIN_URL)
@require_http_methods(['POST'])
def toggle_club_form_channel(request, channel_id, club_id):
    if not is_staff_or_admin(request.user):
        return HttpResponseForbidden("No permission")

    channel = get_object_or_404(FormChannel, pk=channel_id, allow_staff_toggle=True)
    club = get_object_or_404(Club, pk=club_id)
    state, _ = FormChannelClubState.objects.get_or_create(channel=channel, club=club)
    state.is_enabled = not state.is_enabled
    state.updated_by = request.user
    state.save()
    messages.success(request, f"{channel.name} 对 {club.name} 已{'开启' if state.is_enabled else '关闭'}")
    return redirect(request.META.get('HTTP_REFERER') or reverse('clubs:staff_management'))


@login_required(login_url=settings.LOGIN_URL)
@require_http_methods(['GET'])
def zip_download(request):
    submission_key = (request.GET.get('id') or '').strip()
    if not submission_key:
        messages.error(request, '缺少有效的请求编号')
        return redirect(request.META.get('HTTP_REFERER', 'clubs:index'))

    submission = get_object_or_404(
        FormSubmission.objects.select_related('channel', 'club').prefetch_related('uploaded_files__field'),
        public_id=submission_key,
    )
    if not (is_staff_or_admin(request.user) or _is_president_of_club(request.user, submission.club)):
        messages.error(request, '无权下载此提交的材料')
        return redirect('clubs:index')
    if not submission.channel.show_zip_download:
        messages.error(request, '这个通道未开启打包 ZIP 下载')
        return redirect(request.META.get('HTTP_REFERER', 'clubs:index'))

    uploaded_files = [item for item in submission.uploaded_files.all() if item.file]
    if not uploaded_files:
        messages.warning(request, '这个提交没有可打包下载的文件')
        return redirect(request.META.get('HTTP_REFERER', 'clubs:staff_audit_center'))

    import io
    import zipfile
    from pathlib import PurePath

    buffer = io.BytesIO()
    used_names = set()
    with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as archive:
        for uploaded in uploaded_files:
            field_label = uploaded.field.label if uploaded.field_id else '附件'
            filename = uploaded.original_name or PurePath(uploaded.file.name).name
            arcname = f'{field_label}/{filename}'
            base, dot, suffix = arcname.rpartition('.')
            candidate = arcname
            counter = 2
            while candidate in used_names:
                candidate = f'{base}_{counter}.{suffix}' if dot else f'{arcname}_{counter}'
                counter += 1
            used_names.add(candidate)
            uploaded.file.open('rb')
            try:
                archive.writestr(candidate, uploaded.file.read())
            finally:
                uploaded.file.close()

    buffer.seek(0)
    filename = f'{submission.club.name}-{submission.channel.name}-{submission.public_id}.zip'
    response = HttpResponse(buffer.getvalue(), content_type='application/zip')
    response['Content-Disposition'] = f"attachment; filename*=UTF-8''{urllib.parse.quote(filename)}"
    return response




def change_club_status(request, club_id):
    """
    改变社团的活跃状态 - 仅干事可用
    """
    # 检查用户权限 - 仅干事和管理员可以操作
    if not _is_staff(request.user) and not _is_admin(request.user):
        messages.error(request, '您没有权限执行此操作')
        return redirect('clubs:club_detail', club_id=club_id)

    club = get_object_or_404(Club, pk=club_id)

    if request.method == 'POST':
        new_status = request.POST.get('status', club.status)

        # 验证状态值
        valid_statuses = ['active', 'inactive', 'suspended']
        if new_status not in valid_statuses:
            messages.error(request, '无效的社团状态')
            return redirect('clubs:club_detail', club_id=club_id)

        # 如果状态没有改变
        if new_status == club.status:
            messages.info(request, f'社团状态已是 {club.get_status_display()}')
            return redirect('clubs:club_detail', club_id=club_id)

        old_status = club.get_status_display()
        club.status = new_status
        club.save()

        messages.success(request, f'社团状态已从"{old_status}"更改为"{club.get_status_display()}"')
        return redirect('clubs:club_detail', club_id=club_id)

    # GET 请求：返回状态变更选项
    context = {
        'club': club,
        'status_choices': Club._meta.get_field('status').choices,
        'current_status': club.status,
    }
    return render(request, 'clubs/change_club_status.html', context)


@login_required(login_url=settings.LOGIN_URL)
@require_http_methods(["POST"])
def update_club_description(request, club_id):
    club = get_object_or_404(Club, pk=club_id)

    if not (is_staff_or_admin(request.user) or (_is_president(request.user) and club.id in _get_president_club_ids(request.user))):
        return HttpResponseForbidden("您没有权限执行此操作")

    club.description = request.POST.get('description', '').strip()
    club.save(update_fields=['description'])
    messages.success(request, "社团简介已更新")
    return redirect('clubs:club_detail', club_id=club_id)


@login_required(login_url=settings.LOGIN_URL)
@require_http_methods(["GET", "POST"])


@login_required(login_url=settings.LOGIN_URL)
@require_http_methods(["GET", "POST"])
def direct_edit_club_info(request, club_id):
    """
    直接修改社团信息 - 仅干事和管理员可用（无需审核）
    """
    # 检查用户权限 - 仅干事和管理员可以操作
    if not _is_staff(request.user) and not _is_admin(request.user):
        messages.error(request, '您没有权限执行此操作')
        return redirect('clubs:club_detail', club_id=club_id)

    club = get_object_or_404(Club, pk=club_id)

    if request.method == 'POST':
        # 获取表单数据
        new_name = request.POST.get('name', '').strip()
        new_description = request.POST.get('description', '').strip()
        new_founded_date = request.POST.get('founded_date', '')
        new_members_count = request.POST.get('members_count', '')

        # 验证
        errors = []

        if not new_name:
            errors.append('社团名称不能为空')

        if new_members_count and not new_members_count.isdigit():
            errors.append('成员数量必须是数字')

        if errors:
            context = {
                'club': club,
                'errors': errors,
                'form_data': request.POST,
            }
            return render(request, 'clubs/direct_edit_club_info.html', context)

        # 更新社团信息
        if new_name:
            club.name = new_name

        if new_description:
            club.description = new_description
        else:
            club.description = ''

        if new_founded_date:
            club.founded_date = new_founded_date

        if new_members_count:
            club.members_count = int(new_members_count)

        club.save()

        messages.success(request, '社团信息已成功更新！')
        return redirect('clubs:club_detail', club_id=club_id)

    # GET 请求：显示表单
    context = {
        'club': club,
    }
    return render(request, 'clubs/direct_edit_club_info.html', context)


def delete_club(request, club_id):
    """
    删除社团 - 仅干事和管理员可用
    """
    # 检查用户权限 - 仅干事和管理员可以操作
    if not _is_staff(request.user) and not _is_admin(request.user):
        messages.error(request, '您没有权限执行此操作')
        return redirect('clubs:index')

    club = get_object_or_404(Club, pk=club_id)

    if request.method == 'POST':
        # 获取并验证用户输入的确认社团名称
        confirm_club_name = request.POST.get('confirm_club_name')

        # 检查确认社团名称是否正确
        if confirm_club_name == club.name:
                club_name = club.name
                club.delete()
                messages.success(request, f'社团 "{club_name}" 已成功删除')
                return redirect('clubs:index')  # 删除后重定向到首页
        else:
            # 显示错误消息
            messages.error(request, '社团名称输入错误，删除失败！')

    # GET 请求：返回确认页面
    context = {
        'club': club,
    }
    return render(request, 'clubs/delete_club.html', context)


# ==================== 社长换届申请相关视图 ====================






@login_required(login_url=settings.LOGIN_URL)
@require_http_methods(['GET', 'POST'])
def president_member_management(request):
    """社长端社员管理：编辑资料、调整状态、移除社员。"""
    if not _is_president(request.user):
        messages.error(request, '仅社长可访问社员管理')
        return redirect('clubs:index')

    club_ids = _get_president_club_ids(request.user)
    clubs = Club.objects.filter(id__in=club_ids).order_by('name')
    if not clubs.exists():
        messages.error(request, '您当前没有可管理的社团')
        return redirect('clubs:user_dashboard')

    selected_club_id = request.GET.get('club_id') or request.POST.get('club_id')
    if selected_club_id:
        try:
            selected_club_id = int(selected_club_id)
        except ValueError:
            selected_club_id = clubs.first().id
    else:
        selected_club_id = clubs.first().id

    if selected_club_id not in club_ids:
        messages.error(request, '无权限管理该社团社员')
        return redirect('clubs:president_member_management')

    club = get_object_or_404(Club, id=selected_club_id)

    if request.method == 'POST':
        action = request.POST.get('action', '').strip()
        membership_id = request.POST.get('membership_id', '').strip()
        membership = get_object_or_404(ClubMember, id=membership_id, club=club) if membership_id else None

        if action == 'remove_member' and membership:
            if membership.user_profile.role == 'admin':
                messages.error(request, '管理员账号不允许从此处移除')
            else:
                name = membership.user_profile.get_full_name()
                membership.delete()
                messages.success(request, f'已移除社员：{name}')

        elif action == 'set_status' and membership:
            status = request.POST.get('status', 'active').strip()
            if status not in ['active', 'inactive']:
                messages.error(request, '无效的成员状态')
            else:
                membership.status = status
                membership.save(update_fields=['status', 'updated_at'])
                messages.success(request, '成员状态已更新')

        elif action == 'update_profile' and membership:
            profile = membership.user_profile
            email = request.POST.get('email', '').strip()
            real_name = request.POST.get('real_name', '').strip()
            student_id = request.POST.get('student_id', '').strip()
            gender = request.POST.get('gender', '').strip()
            college = request.POST.get('college', '').strip()
            class_name = request.POST.get('class_name', '').strip()
            phone = request.POST.get('phone', '').strip()
            qq = request.POST.get('qq', '').strip()
            wechat = request.POST.get('wechat', '').strip()

            errors = []
            if not real_name:
                errors.append('姓名不能为空')
            if not student_id:
                errors.append('学号不能为空')
            if not phone:
                errors.append('手机号不能为空')
            if not wechat:
                errors.append('微信号不能为空')
            if not email:
                errors.append('邮箱不能为空')

            if student_id and UserProfile.objects.exclude(user=profile.user).filter(student_id=student_id).exists():
                errors.append('学号已被使用')
            if email and User.objects.exclude(id=profile.user_id).filter(email=email).exists():
                errors.append('邮箱已被使用')

            if errors:
                for err in errors:
                    messages.error(request, err)
            else:
                profile.real_name = real_name
                profile.student_id = student_id
                profile.gender = gender
                profile.college = college
                profile.class_name = class_name
                profile.phone = phone
                profile.qq = qq
                profile.wechat = wechat
                profile.save(update_fields=['real_name', 'student_id', 'gender', 'college', 'class_name', 'phone', 'qq', 'wechat', 'updated_at'])

                profile.user.email = email
                profile.user.first_name = real_name
                profile.user.save(update_fields=['email', 'first_name'])
                messages.success(request, '社员信息已更新')

        return redirect(f"{reverse('clubs:president_member_management')}?club_id={club.id}")

    memberships = ClubMember.objects.filter(club=club).select_related('user_profile__user').order_by('status', '-joined_at')

    context = {
        'clubs': clubs,
        'club': club,
        'memberships': memberships,
    }
    return render(request, 'clubs/user/member_management.html', context)





# ==================== 活动申请相关视图 ====================






# ==================== 222房间借用相关视图 ====================

@login_required
def room222_calendar(request):
    """222房间借用日历视图 - 谷歌日历单日网格风格"""
    # 获取日期参数，默认为今天
    date_str = request.GET.get('date')

    if date_str:
        try:
            from datetime import datetime
            view_date = datetime.strptime(date_str, '%Y-%m-%d').date()
        except ValueError:
            view_date = timezone.now().date()
    else:
        view_date = timezone.now().date()

    # 获取前一天和后一天
    from datetime import timedelta
    prev_date = view_date - timedelta(days=1)
    next_date = view_date + timedelta(days=1)

    # 获取当天所有有效的预订
    bookings = Room222Booking.objects.filter(
        booking_date=view_date,
        status='active'
    ).select_related('user__profile', 'club').order_by('start_time')

    # 定义时间段（按照实际课程时间）
    from datetime import time
    raw_time_slots = [
        {'start': time(8, 15), 'end': time(9, 55), 'label': '第1-2节'},
        {'start': time(10, 5), 'end': time(11, 40), 'label': '第3-4节'},
        {'start': time(11, 40), 'end': time(13, 0), 'label': '午休'},
        {'start': time(13, 0), 'end': time(14, 35), 'label': '第5-6节'},
        {'start': time(14, 45), 'end': time(16, 20), 'label': '第7-8节'},
        {'start': time(16, 20), 'end': time(18, 0), 'label': '课外时间'},
        {'start': time(18, 0), 'end': time(19, 0), 'label': '晚餐'},
        {'start': time(19, 0), 'end': time(20, 0), 'label': '晚间1'},
        {'start': time(20, 0), 'end': time(21, 0), 'label': '晚间2'},
        {'start': time(21, 0), 'end': time(22, 0), 'label': '晚间3'},
    ]

    # 计算总分钟数（从8:15到22:00）
    day_start_minutes = 8 * 60 + 15
    total_minutes = (22 * 60) - day_start_minutes  # 13小时45分钟 = 825分钟

    time_slots = []
    for slot in raw_time_slots:
        slot_start_minutes = slot['start'].hour * 60 + slot['start'].minute
        slot_end_minutes = slot['end'].hour * 60 + slot['end'].minute
        duration_minutes = slot_end_minutes - slot_start_minutes
        top_percent = ((slot_start_minutes - day_start_minutes) / total_minutes) * 100
        height_percent = (duration_minutes / total_minutes) * 100

        # 检查该时间段是否有预约覆盖
        has_booking = False
        for booking in bookings:
            booking_start = booking.start_time.hour * 60 + booking.start_time.minute
            booking_end = booking.end_time.hour * 60 + booking.end_time.minute
            # 如果预约时间与时间段有重叠
            if booking_start < slot_end_minutes and booking_end > slot_start_minutes:
                has_booking = True
                break

        time_slots.append({
            'start': slot['start'],
            'end': slot['end'],
            'label': slot['label'],
            'height_percent': height_percent,
            'top_percent': top_percent,
            'has_booking': has_booking,
        })

    # 为每个预订计算其在时间轴上的位置和高度
    bookings_with_position = []
    for booking in bookings:
        # 计算开始时间在时间轴上的位置（以分钟为单位，从8:15开始）
        start_minutes = (booking.start_time.hour * 60 + booking.start_time.minute) - day_start_minutes
        end_minutes = (booking.end_time.hour * 60 + booking.end_time.minute) - day_start_minutes

        # 确保不超出范围
        start_minutes = max(0, start_minutes)
        end_minutes = min(total_minutes, end_minutes)

        # 计算持续时间和位置百分比
        duration_minutes = end_minutes - start_minutes
        top_percent = (start_minutes / total_minutes) * 100
        height_percent = (duration_minutes / total_minutes) * 100

        bookings_with_position.append({
            'booking': booking,
            'top_percent': top_percent,
            'height_percent': height_percent,
            'can_edit': booking.can_edit(request.user),
            'can_delete': booking.can_delete(request.user),
        })

    # 计算周的起始日期（周一）
    from datetime import datetime
    week_start = view_date - timedelta(days=view_date.weekday())
    week_end = week_start + timedelta(days=6)

    context = {
        'view_date': view_date,
        'prev_date': prev_date,
        'next_date': next_date,
        'bookings': bookings_with_position,
        'time_slots': time_slots,
        'is_today': view_date == timezone.now().date(),
        'week_start': week_start,
        'week_end': week_end,
    }
    return render(request, 'clubs/room222_calendar.html', context)


@login_required
def submit_room222_booking(request):
    """提交222房间借用申请 - 无需审核，直接创建"""
    from datetime import datetime
    if request.method == 'POST':
        # 获取表单数据
        club_id = request.POST.get('club_id')
        booking_date_str = request.POST.get('booking_date')
        start_time_str = request.POST.get('start_time')
        end_time_str = request.POST.get('end_time')
        purpose = request.POST.get('purpose')
        participant_count = request.POST.get('participant_count')
        contact_phone = request.POST.get('contact_phone')
        special_requirements = request.POST.get('special_requirements', '')

        # 验证必填字段
        if not all([booking_date_str, start_time_str, end_time_str, purpose,
                   participant_count, contact_phone]):
            messages.error(request, '请填写所有必填字段')
            return redirect('clubs:submit_room222_booking')

        # 转换日期、时间字符串为对象
        from datetime import datetime
        try:
            booking_date = datetime.strptime(booking_date_str, '%Y-%m-%d').date()
            start_time = datetime.strptime(start_time_str, '%H:%M').time()
            end_time = datetime.strptime(end_time_str, '%H:%M').time()
        except ValueError:
            messages.error(request, '时间格式不正确')
            return redirect('clubs:submit_room222_booking')

        # 获取社团（如果选择了）
        club = None
        if club_id:
            club = get_object_or_404(Club, pk=club_id)
            # 验证用户是否为该社团社长
            if club.president != request.user:
                messages.error(request, '您不是该社团的社长')
                return redirect('clubs:submit_room222_booking')

        # 创建借用记录（状态为active，无需审核）
        booking = Room222Booking(
            user=request.user,
            club=club,
            booking_date=booking_date,
            start_time=start_time,
            end_time=end_time,
            purpose=purpose,
            participant_count=participant_count,
            contact_phone=contact_phone,
            special_requirements=special_requirements,
            status='active'
        )

        # 检查时间冲突
        if booking.has_conflict():
            messages.error(request, '该时间段已被预订，请选择其他时间')
            return redirect('clubs:submit_room222_booking')

        booking.save()
        messages.success(request, '222房间预约成功！')
        return redirect('clubs:room222_calendar')

    # GET 请求
    # 获取用户是社长的所有社团
    user_clubs = Club.objects.filter(
        officers__user_profile__user=request.user,
        officers__position='president',
        officers__is_current=True,
        status='active',
    )

    # 获取日期参数（从日历页面跳转过来时）
    selected_date = request.GET.get('date', '')
    selected_start_time = request.GET.get('start_time', '')
    selected_end_time = request.GET.get('end_time', '')

    # 获取今天的日期
    today = timezone.now().date().strftime('%Y-%m-%d')

    # 获取当天已存在的预约，用于前端快速校验
    booked_intervals = []
    if selected_date:
        try:
            selected_date_obj = datetime.strptime(selected_date, '%Y-%m-%d').date()
            existing = Room222Booking.objects.filter(
                booking_date=selected_date_obj,
                status='active'
            ).values('start_time', 'end_time')
            booked_intervals = [
                {
                    'start': item['start_time'].strftime('%H:%M'),
                    'end': item['end_time'].strftime('%H:%M'),
                }
                for item in existing
            ]
        except ValueError:
            booked_intervals = []

    context = {
        'user_clubs': user_clubs,
        'selected_date': selected_date,
        'selected_start_time': selected_start_time,
        'selected_end_time': selected_end_time,
        'today': today,
        'booked_intervals': booked_intervals,
    }
    return render(request, 'clubs/submit_room222_booking.html', context)


@login_required
def my_room222_bookings(request):
    """查看我的222房间借用记录"""
    bookings = Room222Booking.objects.filter(
        user=request.user
    ).order_by('-booking_date', '-start_time')

    context = {
        'bookings': bookings,
    }
    return render(request, 'clubs/my_room222_bookings.html', context)


@login_required
def edit_room222_booking(request, booking_id):
    """编辑222房间借用"""
    booking = get_object_or_404(Room222Booking, pk=booking_id)

    # 检查权限
    if not booking.can_edit(request.user):
        messages.error(request, '您没有权限编辑此预约')
        return redirect('clubs:my_room222_bookings')

    if request.method == 'POST':
        # 获取表单数据
        club_id = request.POST.get('club_id')
        booking_date = request.POST.get('booking_date')
        start_time = request.POST.get('start_time')
        end_time = request.POST.get('end_time')
        purpose = request.POST.get('purpose')
        participant_count = request.POST.get('participant_count')
        contact_phone = request.POST.get('contact_phone')
        special_requirements = request.POST.get('special_requirements', '')

        # 验证必填字段
        if not all([booking_date, start_time, end_time, purpose,
                   participant_count, contact_phone]):
            messages.error(request, '请填写所有必填字段')
            return redirect('clubs:edit_room222_booking', booking_id=booking_id)

        # 获取社团（如果选择了）
        club = None
        if club_id:
            club = get_object_or_404(Club, pk=club_id)

        # 更新预订信息
        booking.club = club
        booking.booking_date = booking_date
        booking.start_time = start_time
        booking.end_time = end_time
        booking.purpose = purpose
        booking.participant_count = participant_count
        booking.contact_phone = contact_phone
        booking.special_requirements = special_requirements

        # 检查时间冲突
        if booking.has_conflict():
            messages.error(request, '该时间段已被其他预订占用，请选择其他时间')
            return redirect('clubs:edit_room222_booking', booking_id=booking_id)

        booking.save()
        messages.success(request, '预约已成功更新')
        return redirect('clubs:my_room222_bookings')

    # GET 请求
    user_clubs = Club.objects.filter(
        officers__user_profile__user=request.user,
        officers__position='president',
        officers__is_current=True,
        status='active',
    )

    context = {
        'booking': booking,
        'user_clubs': user_clubs,
    }
    return render(request, 'clubs/edit_room222_booking.html', context)


@login_required
def delete_room222_booking(request, booking_id):
    """删除/取消222房间借用"""
    booking = get_object_or_404(Room222Booking, pk=booking_id)

    # 检查权限
    if not booking.can_delete(request.user):
        messages.error(request, '您没有权限删除此预约')
        return redirect('clubs:my_room222_bookings')

    if request.method == 'POST':
        booking_info = f"{booking.booking_date} {booking.start_time}-{booking.end_time}"
        booking.status = 'cancelled'
        booking.save()
        messages.success(request, f'已取消预约：{booking_info}')
        return redirect('clubs:my_room222_bookings')

    context = {
        'booking': booking,
    }
    return render(request, 'clubs/delete_room222_booking.html', context)


# ==================== 活动申请功能 ====================





# ==================== 审核活动申请和社长换届 ====================
























AUDIT_REQUEST_MODELS = {}


def _audit_tab_url(tab):
    return tab.replace('_', '-')







@login_required
@require_http_methods(["GET", "POST"])
def edit_department_introduction(request):
    """编辑部门介绍页面 - 仅管理员可访问"""
    # 检查权限 - 仅管理员可编辑
    if not _is_admin(request.user) and not request.user.is_superuser:
        messages.error(request, '您没有权限编辑部门介绍')
        return redirect('clubs:index')

    # 获取所有部门
    departments = DepartmentIntroduction.objects.all().order_by('department')

    if request.method == 'POST':
        # 处理表单提交 - 更新每个部门的介绍
        try:
            for dept in departments:
                # 获取表单数据
                dept_key = f"dept_{dept.id}"
                description = request.POST.get(f"{dept_key}_description", "").strip()
                highlights = request.POST.get(f"{dept_key}_highlights", "").strip()
                icon = request.POST.get(f"{dept_key}_icon", dept.icon).strip()

                # 更新部门信息
                if description:  # 只有在有描述时才更新
                    dept.description = description
                    dept.highlights = highlights
                    dept.icon = icon
                    dept.updated_by = request.user
                    dept.save()
                    messages.success(request, f'已更新{dept.get_department_display()}部门信息')

            messages.success(request, '所有部门介绍已成功更新')
            return redirect('clubs:index')
        except Exception as e:
            messages.error(request, f'更新部门介绍时出错: {str(e)}')

    context = {
        'departments': departments,
        'material_icons': [
            'work', 'assessment', 'event', 'people', 'speaker',
            'star', 'explore', 'check', 'info', 'favorite',
            'school', 'business', 'settings', 'security', 'trending_up'
        ]
    }
    return render(request, 'clubs/edit_department_introduction.html', context)


# ==================== 材料要求管理视图 ====================









@login_required
def manage_departments(request):
    """管理部门"""
    if not is_staff_or_admin(request.user):
        messages.error(request, '您没有权限执行此操作')
        return redirect('clubs:index')

    departments = Department.objects.all()
    return render(request, 'clubs/admin/manage_departments.html', {'departments': departments})


@login_required
def add_department(request):
    """添加部门"""
    if not is_staff_or_admin(request.user):
        messages.error(request, '您没有权限执行此操作')
        return redirect('clubs:index')

    if request.method == 'POST':
        name = request.POST.get('name')
        description = request.POST.get('description')
        highlights = request.POST.get('highlights')
        icon = request.POST.get('icon', 'work')
        order = request.POST.get('order', 0)

        try:
            Department.objects.create(
                name=name,
                description=description,
                highlights=highlights,
                icon=icon,
                order=int(order)
            )
            messages.success(request, '部门添加成功')
            return redirect('clubs:manage_departments')
        except Exception as e:
            messages.error(request, f'添加失败: {str(e)}')

    return render(request, 'clubs/admin/department_form.html', {'title': '添加部门'})


@login_required
def get_clubs_list(request):
    """获取社团列表API"""
    if not is_staff_or_admin(request.user):
        return HttpResponseForbidden()

    clubs = Club.objects.filter(status='active').values('id', 'name')
    return JsonResponse({'clubs': list(clubs)})




@login_required
def room_calendar(request):
    """
    显示房间预约日历
    """
    # 获取参数
    room_id = request.GET.get('room_id')
    date_str = request.GET.get('date')

    # 获取当前查看的日期
    if date_str:
        try:
            view_date = datetime.strptime(date_str, '%Y-%m-%d').date()
        except ValueError:
            view_date = timezone.now().date()
    else:
        view_date = timezone.now().date()

    # 获取所有可用房间
    rooms = Room.objects.filter(status='available')
    if not rooms.exists():
        messages.error(request, '暂时没有可用的房间')
        return redirect('clubs:index')

    # 确定当前选中的房间
    if room_id:
        selected_room = get_object_or_404(Room, pk=room_id)
    else:
        selected_room = rooms.first()

    # 计算日期导航
    is_today = (view_date == timezone.now().date())
    prev_date = view_date - timezone.timedelta(days=1)
    next_date = view_date + timezone.timedelta(days=1)
    week_start = view_date - timezone.timedelta(days=view_date.weekday())

    # 获取时间段
    time_slots = TimeSlot.objects.filter(is_active=True).order_by('start_time')

    # 计算整个日历的起止时间（用于计算百分比）
    # 默认 8:00 (480min) 到 22:00 (1320min)，总长 840min
    day_start_minutes = 8 * 60
    day_end_minutes = 22 * 60

    if time_slots.exists():
        first_slot = time_slots.first()
        last_slot = time_slots.last()
        day_start_minutes = min(day_start_minutes, first_slot.start_time.hour * 60 + first_slot.start_time.minute)
        day_end_minutes = max(day_end_minutes, last_slot.end_time.hour * 60 + last_slot.end_time.minute)

    total_minutes = day_end_minutes - day_start_minutes
    if total_minutes <= 0:
        total_minutes = 14 * 60

    processed_slots = []
    for slot in time_slots:
        slot_start_min = slot.start_time.hour * 60 + slot.start_time.minute
        slot_end_min = slot.end_time.hour * 60 + slot.end_time.minute

        # 使用浮点数计算百分比，保留4位小数以确保精度
        top_percent = ((slot_start_min - day_start_minutes) / total_minutes) * 100
        height_percent = ((slot_end_min - slot_start_min) / total_minutes) * 100

        # 检查该时间段是否已有预约
        has_booking = RoomBooking.objects.filter(
            room=selected_room,
            booking_date=view_date,
            status='active',
            start_time__lt=slot.end_time,
            end_time__gt=slot.start_time
        ).exists()

        processed_slots.append({
            'start': slot.start_time,
            'end': slot.end_time,
            'label': slot.label,
            'top_percent': f"{top_percent:.4f}",  # 格式化为字符串，避免本地化问题
            'height_percent': f"{height_percent:.4f}", # 格式化为字符串
            'has_booking': has_booking,
            'id': slot.id
        })

    # 获取当天的预约
    bookings = RoomBooking.objects.filter(
        room=selected_room,
        booking_date=view_date,
        status='active'
    ).select_related('user__profile', 'club')

    processed_bookings = []
    for booking in bookings:
        # 权限检查
        can_edit = booking.can_edit(request.user)
        can_delete = booking.can_delete(request.user)

        # 计算位置
        b_start_min = booking.start_time.hour * 60 + booking.start_time.minute
        b_end_min = booking.end_time.hour * 60 + booking.end_time.minute

        top_percent = ((b_start_min - day_start_minutes) / total_minutes) * 100
        height_percent = ((b_end_min - b_start_min) / total_minutes) * 100

        processed_bookings.append({
            'booking': booking,
            'top_percent': f"{top_percent:.4f}", # 格式化为字符串
            'height_percent': f"{height_percent:.4f}", # 格式化为字符串
            'can_edit': can_edit,
            'can_delete': can_delete
        })

    context = {
        'rooms': rooms,
        'selected_room': selected_room,
        'view_date': view_date,
        'is_today': is_today,
        'prev_date': prev_date,
        'next_date': next_date,
        'week_start': week_start,
        'time_slots': processed_slots,
        'bookings': processed_bookings,
    }

    return render(request, 'clubs/room_calendar.html', context)



@login_required
def edit_department(request, dept_id):
    """编辑部门"""
    if not is_staff_or_admin(request.user):
        messages.error(request, '您没有权限执行此操作')
        return redirect('clubs:index')

    dept = get_object_or_404(Department, pk=dept_id)

    if request.method == 'POST':
        dept.name = request.POST.get('name')
        dept.description = request.POST.get('description')
        dept.highlights = request.POST.get('highlights')
        dept.icon = request.POST.get('icon', 'work')
        dept.order = int(request.POST.get('order', 0))
        dept.updated_by = request.user
        dept.save()
        messages.success(request, '部门更新成功')
        return redirect('clubs:manage_departments')

    return render(request, 'clubs/admin/department_form.html', {
        'title': '编辑部门',
        'department': dept
    })


@login_required
def delete_department(request, dept_id):
    """删除部门"""
    if not is_staff_or_admin(request.user):
        messages.error(request, '您没有权限执行此操作')
        return redirect('clubs:index')

    dept = get_object_or_404(Department, pk=dept_id)
    if request.method == 'POST':
        dept.delete()
        messages.success(request, '部门删除成功')

    return redirect('clubs:manage_departments')


@login_required
@login_required
@require_http_methods(["GET", "POST"])
def submit_room_booking(request):
    """提交房间预约"""
    if request.method == 'GET':
        # 获取预填参数
        room_id = request.GET.get('room_id')
        date_str = request.GET.get('date')
        start_time = request.GET.get('start_time')
        end_time = request.GET.get('end_time')

        # 获取基础数据
        rooms = Room.objects.filter(status='available')

        # 获取用户关联的社团（作为社长，通过 Officer 表）
        user_clubs = []
        if hasattr(request.user, 'profile'):
            user_clubs = Club.objects.filter(
                officers__user_profile__user=request.user,
                officers__position='president',
                officers__is_current=True,
                status='active',
            )

        today = timezone.now().date().strftime('%Y-%m-%d')

        context = {
            'rooms': rooms,
            'user_clubs': user_clubs,
            'today': today,
            'selected_room_id': int(room_id) if room_id and room_id.isdigit() else None,
            'selected_date': date_str,
            'selected_start_time': start_time,
            'selected_end_time': end_time,
            'is_staff_or_admin': is_staff_or_admin(request.user)
        }
        return render(request, 'clubs/submit_room_booking.html', context)

    # POST 请求处理
    try:
        room_id = request.POST.get('room_id')
        date_str = request.POST.get('booking_date')
        start_time_str = request.POST.get('start_time')
        end_time_str = request.POST.get('end_time')
        club_id = request.POST.get('club_id')
        purpose = request.POST.get('purpose')
        contact_phone = request.POST.get('contact_phone')
        special_requirements = request.POST.get('special_requirements')
        participant_count = request.POST.get('participant_count')

        if not all([room_id, date_str, start_time_str, end_time_str, purpose, contact_phone, participant_count]):
            messages.error(request, '请填写所有必填项')
            return redirect('clubs:submit_room_booking')

        # 验证手机号格式
        if not re.match(r'^1[3-9]\d{9}$', contact_phone):
            messages.error(request, '请输入有效的11位手机号码')
            return redirect('clubs:submit_room_booking')

        room = get_object_or_404(Room, pk=room_id)
        booking_date = datetime.strptime(date_str, '%Y-%m-%d').date()

        # 解析时间
        start_time = datetime.strptime(start_time_str, '%H:%M').time()
        end_time = datetime.strptime(end_time_str, '%H:%M').time()

        # 验证时间顺序
        if start_time >= end_time:
            messages.error(request, '结束时间必须晚于开始时间')
            return redirect('clubs:submit_room_booking')

        # 验证是否为有效的固定时间段
        is_valid_slot = TimeSlot.objects.filter(
            start_time=start_time,
            end_time=end_time,
            is_active=True
        ).exists()

        if not is_valid_slot:
            messages.error(request, '请选择有效的固定时间段')
            return redirect('clubs:room_calendar')

        # 确定社团
        club = None
        if club_id:
            club = get_object_or_404(Club, pk=club_id)
            # 验证用户是否有权代表该社团申请（如果是社长）
            if not is_staff_or_admin(request.user):
                if club.president != request.user:
                    messages.error(request, '您不是该社团的社长，无法代表申请')
                    return redirect('clubs:submit_room_booking')
        else:
            # 个人申请，必须是干事或管理员
            if not is_staff_or_admin(request.user):
                messages.error(request, '普通用户必须选择社团进行申请')
                return redirect('clubs:submit_room_booking')

        # 检查冲突
        existing_booking = RoomBooking.objects.filter(
            room=room,
            booking_date=booking_date,
            status='active'
        ).filter(
            Q(start_time__lt=end_time) &
            Q(end_time__gt=start_time)
        ).exists()

        if existing_booking:
            messages.error(request, '该时间段已被预约，请选择其他时间')
            # 返回并带上参数以便重填，这里简单处理直接跳回日历
            return redirect('clubs:room_calendar')

        # 创建预约
        RoomBooking.objects.create(
            room=room,
            user=request.user,
            club=club,
            booking_date=booking_date,
            start_time=start_time,
            end_time=end_time,
            purpose=purpose,
            contact_phone=contact_phone,
            special_requirements=special_requirements,
            participant_count=int(participant_count),
            status='active'
        )
        messages.success(request, '预约提交成功')
    except Exception as e:
        messages.error(request, f'预约失败: {str(e)}')
        # 发生错误时，最好能保留用户输入，但这里简化处理
        return redirect('clubs:room_calendar')

    return redirect('clubs:room_calendar')


@login_required
def my_room_bookings(request):
    """我的预约"""
    bookings = RoomBooking.objects.filter(user=request.user).order_by('-booking_date', '-start_time')
    return render(request, 'clubs/room_my_bookings.html', {'bookings': bookings})


@login_required
def edit_room_booking(request, booking_id):
    """编辑预约"""
    booking = get_object_or_404(RoomBooking, pk=booking_id)
    if not booking.can_edit(request.user):
        messages.error(request, '您没有权限编辑此预约')
        return redirect('clubs:my_room_bookings')

    if request.method == 'POST':
        # 简单实现，实际可能需要更多逻辑
        booking.reason = request.POST.get('reason')
        booking.save()
        messages.success(request, '预约已更新')
        return redirect('clubs:my_room_bookings')

    return render(request, 'clubs/room_my_bookings.html', {'booking': booking})


@login_required
def delete_room_booking(request, booking_id):
    """取消预约"""
    booking = get_object_or_404(RoomBooking, pk=booking_id)
    if not booking.can_delete(request.user):
        messages.error(request, '您没有权限取消此预约')
    else:
        booking.delete()
        messages.success(request, '预约已取消')
    return redirect('clubs:my_room_bookings')


@login_required
def admin_room_add(request):
    """管理员-添加房间"""
    if not is_staff_or_admin(request.user):
        return redirect('clubs:index')
    if request.method == 'POST':
        Room.objects.create(
            name=request.POST.get('name'),
            capacity=request.POST.get('capacity'),
            location=request.POST.get('location'),
            description=request.POST.get('description'),
            status=request.POST.get('status')
        )
        return redirect('clubs:admin_booking_management')
    return render(request, 'clubs/admin/room_form.html')


@login_required
def admin_room_edit(request, room_id):
    """管理员-编辑房间"""
    if not is_staff_or_admin(request.user):
        return redirect('clubs:index')
    room = get_object_or_404(Room, pk=room_id)
    if request.method == 'POST':
        room.name = request.POST.get('name')
        room.capacity = request.POST.get('capacity')
        room.location = request.POST.get('location')
        room.description = request.POST.get('description')
        room.status = request.POST.get('status')
        room.save()
        return redirect('clubs:admin_booking_management')
    return render(request, 'clubs/admin/room_form.html', {'room': room})


@login_required
def admin_room_delete(request, room_id):
    """管理员-删除房间"""
    if not is_staff_or_admin(request.user):
        return redirect('clubs:index')
    room = get_object_or_404(Room, pk=room_id)
    if request.method == 'POST':
        room.delete()
    return redirect('clubs:admin_booking_management')


@login_required
def admin_booking_management(request):
    """管理员-预约管理"""
    if not is_staff_or_admin(request.user):
        return redirect('clubs:index')
    bookings = RoomBooking.objects.select_related('room', 'user', 'club').order_by('-booking_date', '-start_time')
    rooms = Room.objects.all().order_by('name')
    slots = TimeSlot.objects.all().order_by('start_time')
    return render(
        request,
        'clubs/admin/booking_management.html',
        {
            'bookings': bookings,
            'rooms': rooms,
            'time_slots': slots,
        },
    )


@login_required
def admin_time_slot_add(request):
    """管理员-添加时间段"""
    if not is_staff_or_admin(request.user):
        return redirect('clubs:index')
    if request.method == 'POST':
        TimeSlot.objects.create(
            start_time=request.POST.get('start_time'),
            end_time=request.POST.get('end_time'),
            label=request.POST.get('label'),
            is_active=request.POST.get('is_active') == 'on'
        )
        return redirect('clubs:admin_booking_management')
    return render(request, 'clubs/admin/time_slot_form.html')


@login_required
def admin_time_slot_edit(request, slot_id):
    """管理员-编辑时间段"""
    if not is_staff_or_admin(request.user):
        return redirect('clubs:index')
    slot = get_object_or_404(TimeSlot, pk=slot_id)
    if request.method == 'POST':
        slot.start_time = request.POST.get('start_time')
        slot.end_time = request.POST.get('end_time')
        slot.label = request.POST.get('label')
        slot.is_active = request.POST.get('is_active') == 'on'
        slot.save()
        return redirect('clubs:admin_booking_management')
    return render(request, 'clubs/admin/time_slot_form.html', {'slot': slot})


@login_required
def admin_time_slot_delete(request, slot_id):
    """管理员-删除时间段"""
    if not is_staff_or_admin(request.user):
        return redirect('clubs:index')
    slot = get_object_or_404(TimeSlot, pk=slot_id)
    if request.method == 'POST':
        slot.delete()
    return redirect('clubs:admin_booking_management')


@login_required(login_url=settings.LOGIN_URL)
def manage_favicon(request):
    """管理网站图标 + 字体设置"""
    if not is_staff_or_admin(request.user):
        messages.error(request, '权限不足')
        return redirect('clubs:index')

    if request.method == 'POST':
        form_type = request.POST.get('form_type', 'favicon')
        if form_type == 'font_settings':
            cfg = SiteSettings.get_settings()
            cfg.font_icon_url = (
                request.POST.get('font_icon_url', '').strip()
                or 'https://fonts.font.im/icon?family=Material+Icons'
            )
            cfg.body_font_url = request.POST.get('body_font_url', '').strip()
            cfg.body_font_family = request.POST.get('body_font_family', '').strip()
            cfg.save()
            messages.success(request, '站点字体设置已保存，刷新页面后生效')
            return redirect('clubs:manage_favicon')
        elif 'favicon' in request.FILES:
            upload = request.FILES['favicon']
            ok, logo_message = process_site_logo(upload, allow_webp=False)
            if ok:
                messages.success(request, logo_message)
            else:
                messages.error(request, logo_message)
            return redirect('clubs:manage_favicon')

    cfg = SiteSettings.get_settings()
    presets = [
        {'label': '默认镜像 (fonts.font.im)',    'value': 'https://fonts.font.im/icon?family=Material+Icons'},
        {'label': 'fonts.googleapis.com（需翻墙）', 'value': 'https://fonts.googleapis.com/icon?family=Material+Icons'},
        {'label': 'SJTUG 镜像',                 'value': 'https://google-fonts.mirrors.sjtug.sjtu.edu.cn/icon?family=Material+Icons'},
        {'label': '中科大镜像',                  'value': 'https://fonts.loli.net/icon?family=Material+Icons'},
    ]
    body_font_presets = [
        {'label': '不使用外部字体',                   'url': '', 'family': ''},
        {'label': 'Noto Sans SC (fonts.font.im)', 'url': 'https://fonts.font.im/css2?family=Noto+Sans+SC:wght@400;500;700&display=swap', 'family': "'Noto Sans SC', sans-serif"},
        {'label': 'Noto Sans SC (SJTUG)',         'url': 'https://google-fonts.mirrors.sjtug.sjtu.edu.cn/css2?family=Noto+Sans+SC:wght@400;500;700&display=swap', 'family': "'Noto Sans SC', sans-serif"},
        {'label': 'Noto Sans SC (loli.net)',      'url': 'https://fonts.loli.net/css2?family=Noto+Sans+SC:wght@400;500;700&display=swap', 'family': "'Noto Sans SC', sans-serif"},
    ]
    icon_preview_list = [
        'home', 'settings', 'people', 'notifications', 'search', 'dashboard',
        'add_circle', 'edit', 'delete', 'check_circle', 'cancel', 'upload_file',
        'arrow_back', 'save', 'font_download',
    ]
    return render(request, 'clubs/admin/manage_favicon.html', {
        'cfg': cfg,
        'presets': presets,
        'body_font_presets': body_font_presets,
        'icon_preview_list': icon_preview_list,
    })


@login_required(login_url='clubs:login')
def get_department_members(request, department_id):
    """API: 获取部门成员列表"""
    try:
        department = Department.objects.get(id=department_id)
    except Department.DoesNotExist:
        return JsonResponse({'error': '部门不存在'}, status=404)

    # Include both staff and admin users in the department
    members = UserProfile.objects.filter(
        role__in=['staff', 'admin'],
        department_link=department
    ).select_related('user').order_by('staff_level', 'user__username')

    directors_data = []
    members_data = []

    for member in members:
        avatar_url = member.avatar.url if member.avatar else None
        item = {
            'id': member.user.id,
            'name': member.get_full_name(),
            'avatar': avatar_url,
            'initial': member.get_full_name()[0].upper() if member.get_full_name() else member.user.username[0].upper(),
        }
        if member.staff_level == 'director':
            directors_data.append(item)
        else:
            members_data.append(item)

    return JsonResponse({
        'name': department.name,
        'directors': directors_data,
        'members': members_data
    })


@login_required(login_url=settings.LOGIN_URL)
def admin_assign_presidents(request):
    """管理员批量指定/修改各社团的现任社长。

    GET  — 列出所有社团及其当前绑定的社长，每行有下拉框选择用户。
    POST — 根据提交数据，在 Officer 表中更新或创建社长记录。
    """
    if not _is_admin(request.user):
        messages.error(request, '仅管理员可以访问此页面')
        return redirect('clubs:index')

    if request.method == 'POST':
        updated = 0
        cleared = 0
        today = timezone.now().date()

        clubs_all = Club.objects.all()
        for club in clubs_all:
            field_key = f'president_{club.id}'
            user_id_str = request.POST.get(field_key, '').strip()

            if user_id_str:
                try:
                    new_user = User.objects.get(pk=int(user_id_str))
                    new_profile = new_user.profile  # type: ignore[attr-defined]
                except (User.DoesNotExist, UserProfile.DoesNotExist, ValueError):
                    continue

                # 将该社团所有旧社长标记为离任
                Officer.objects.filter(
                    club=club, position='president', is_current=True
                ).exclude(user_profile=new_profile).update(
                    is_current=False, end_date=today
                )

                # 创建或激活该社长的 Officer 记录
                obj, created = Officer.objects.get_or_create(
                    club=club,
                    user_profile=new_profile,
                    position='president',
                    defaults={'appointed_date': today, 'is_current': True},
                )
                if not created and not obj.is_current:
                    obj.is_current = True
                    obj.end_date = None
                    obj.appointed_date = today
                    obj.save(update_fields=['is_current', 'end_date', 'appointed_date'])
                updated += 1
            else:
                # 用户选择了"无社长"——将现有社长全部标记离任
                count = Officer.objects.filter(
                    club=club, position='president', is_current=True
                ).update(is_current=False, end_date=today)
                if count:
                    cleared += count

        messages.success(request, f'已更新 {updated} 个社团的社长绑定，清除 {cleared} 个离任记录。')
        return redirect('clubs:admin_assign_presidents')

    # GET — 准备数据
    clubs_qs = Club.objects.prefetch_related(
        Prefetch(
            'officers',
            queryset=Officer.objects.filter(position='president', is_current=True)
                .select_related('user_profile__user'),
            to_attr='_president_list',
        )
    ).order_by('name')

    # 所有 president 角色用户，用于下拉框
    president_users = User.objects.filter(profile__role='president').select_related('profile').order_by('profile__real_name')

    clubs_data = []
    for club in clubs_qs:
        pres_list = getattr(club, '_president_list', [])
        current_president_user = (
            pres_list[0].user_profile.user if pres_list and pres_list[0].user_profile else None
        )
        clubs_data.append({
            'club': club,
            'current_president': current_president_user,
        })

    assigned_count = sum(1 for c in clubs_data if c['current_president'])
    context = {
        'clubs_data': clubs_data,
        'president_users': president_users,
        'total_clubs': len(clubs_data),
        'assigned_count': assigned_count,
        'unassigned_count': len(clubs_data) - assigned_count,
    }
    return render(request, 'clubs/admin/assign_presidents.html', context)


# ==================== Dynamic form system ====================

def _active_channels():
    return (
        FormChannel.objects.filter(is_active=True)
        .exclude(slug='')
        .order_by('order', 'id')
    )


def _get_channel(slug, active_only=True):
    qs = FormChannel.objects.all()
    if active_only:
        qs = qs.filter(is_active=True)
    return get_object_or_404(qs, slug=slug)


def _channel_enabled_for_club(channel, club):
    state = FormChannelClubState.objects.filter(channel=channel, club=club).first()
    return True if state is None else state.is_enabled


def _current_cycle(channel):
    if channel.submission_policy != 'once_per_cycle':
        return None
    return channel.cycles.filter(is_active=True).order_by('-sequence', '-starts_at').first()


def _submission_policy_scope(channel, club, user=None):
    cycle = _current_cycle(channel)
    qs = FormSubmission.objects.filter(channel=channel, club=club)
    if user is not None:
        qs = qs.filter(submitter=user)
    if channel.submission_policy == 'once_per_cycle':
        if not cycle:
            return None, qs.none(), '当前通道没有开启的周期'
        qs = qs.filter(cycle=cycle)
    return cycle, qs, ''


def _can_create_submission(channel, club, user):
    if not _channel_enabled_for_club(channel, club):
        return False, None, '该社团暂未开放此提交通道'
    cycle, scoped, error = _submission_policy_scope(channel, club, user)
    if error:
        return False, cycle, error
    if channel.submission_policy in ['once_total', 'once_per_cycle']:
        blocking = scoped.filter(status__in=['pending', 'approved']).order_by('-submitted_at').first()
        if blocking:
            return False, cycle, f'该通道已有{blocking.get_status_display()}提交，不能重复提交'
    return True, cycle, ''


def _field_input_name(field):
    return f'field_{field.id}'


def _get_submission_or_404(submission_key):
    return get_object_or_404(
        FormSubmission.objects.select_related('channel', 'club', 'submitter', 'reviewer'),
        public_id=submission_key,
    )


def _safe_filename_part(value, fallback='文件'):
    value = (value or fallback).strip() or fallback
    value = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', '_', value)
    value = re.sub(r'\s+', '', value)
    return value[:80] or fallback


def _renamed_upload_name(field, uploaded, index=1, total=1):
    ext = os.path.splitext(uploaded.name)[1].lower()
    base = _safe_filename_part(field.label, '附件')
    suffix = str(index) if total > 1 else ''
    return f'{base}{suffix}{ext}'


def _delete_uploaded_record(uploaded):
    if uploaded.file:
        uploaded.file.delete(save=False)
    uploaded.delete()


def _parse_options(raw):
    return [line.strip() for line in (raw or '').replace('\r', '').split('\n') if line.strip()]


def _extensions_support_word_merge(extensions):
    normalized = {str(item).strip().lower() for item in extensions if str(item).strip()}
    return bool(normalized) and normalized.issubset(FormField.WORD_MERGE_EXTENSIONS)


def _parse_validation(field_type, post):
    validation = {}
    if field_type == 'file':
        exts = post.get('allowed_extensions', '').strip()
        validation['allowed_extensions'] = [e.strip() if e.strip().startswith('.') else f'.{e.strip()}' for e in exts.split(',') if e.strip()] or ['.doc', '.docx', '.pdf', '.jpg', '.jpeg', '.png', '.zip']
        validation['allow_multiple'] = post.get('allow_multiple') == 'on'
        validation['merge_to_word'] = (
            validation['allow_multiple']
            and post.get('merge_to_word') == 'on'
            and _extensions_support_word_merge(validation['allowed_extensions'])
        )
        try:
            validation['max_size_mb'] = int(post.get('max_size_mb', '10') or 10)
        except ValueError:
            validation['max_size_mb'] = 10
    elif field_type == 'number':
        for key in ['min', 'max', 'step']:
            value = post.get(key, '').strip()
            if value:
                validation[key] = value
    else:
        max_length = post.get('max_length', '').strip()
        if max_length:
            validation['max_length'] = max_length
    return validation


def _is_president_of_club(user, club):
    return Officer.objects.filter(
        club=club,
        user_profile__user=user,
        position='president',
        is_current=True,
    ).exists()


def _coerce_bool_text(value):
    return str(value).strip() in ['是', 'true', 'True', '1', 'yes', '公开']


def _validate_dynamic_submission(channel, post, files, fields=None, force_required_field_ids=None):
    cleaned = {}
    upload_map = {}
    errors = []
    fields = fields or channel.fields.filter(is_active=True).order_by('order', 'id')
    force_required_field_ids = set(force_required_field_ids or [])
    for field in fields:
        name = _field_input_name(field)
        is_required = field.required or field.id in force_required_field_ids
        if field.field_type == 'file':
            uploaded_files = files.getlist(name)
            if is_required and not uploaded_files:
                errors.append(f'{field.label} 为必填文件')
                continue
            if uploaded_files and not field.allow_multiple_files() and len(uploaded_files) > 1:
                errors.append(f'{field.label} 只能上传 1 个文件')
                continue
            valid_uploads = []
            for uploaded in uploaded_files:
                ext = os.path.splitext(uploaded.name)[1].lower()
                allowed = field.allowed_extensions()
                if ext not in allowed:
                    errors.append(f'{field.label} 文件类型不允许，允许：{", ".join(allowed)}')
                max_bytes = field.max_size_mb() * 1024 * 1024
                if uploaded.size > max_bytes:
                    errors.append(f'{field.label} 文件不能超过 {field.max_size_mb()}MB')
                valid_uploads.append(uploaded)
            if valid_uploads:
                upload_map[field.id] = valid_uploads
            continue

        if field.field_type == 'checkbox':
            value = post.getlist(name)
        else:
            value = post.get(name, '').strip()

        if is_required and (value == '' or value == []):
            errors.append(f'{field.label} 为必填项')
            continue

        if field.field_type == 'number' and value != '':
            try:
                float(value)
            except ValueError:
                errors.append(f'{field.label} 必须是数字')
        max_length = field.validation.get('max_length')
        if max_length and isinstance(value, str) and len(value) > int(max_length):
            errors.append(f'{field.label} 不能超过 {max_length} 个字符')
        if field.field_type in ['select', 'radio'] and value and field.options and value not in field.options:
            errors.append(f'{field.label} 选项无效')
        if field.field_type == 'checkbox' and value and field.options:
            invalid = [item for item in value if item not in field.options]
            if invalid:
                errors.append(f'{field.label} 选项无效')
        cleaned[field.id] = value
    return fields, cleaned, upload_map, errors


def _add_image_to_document(document, image_source):
    from io import BytesIO
    from docx.shared import Inches
    from PIL import Image

    if isinstance(image_source, (str, os.PathLike)):
        image_buffer = BytesIO()
        with Image.open(image_source) as image:
            if image.mode not in ['RGB', 'RGBA']:
                image = image.convert('RGB')
            image.save(image_buffer, format='PNG')
        image_buffer.seek(0)
        document.add_picture(image_buffer, width=Inches(6))
        return

    document.add_picture(image_source, width=Inches(6))


def _append_docx_text(document, path):
    from docx import Document

    source = Document(path)
    appended = False
    for paragraph in source.paragraphs:
        text = paragraph.text.strip()
        if text:
            document.add_paragraph(text)
            appended = True
    for table in source.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                document.add_paragraph(' | '.join(cells))
                appended = True
    if not appended:
        document.add_paragraph('这个 Word 文档没有可提取的正文内容。')


def _render_pdf_pages_to_document(document, path):
    from io import BytesIO
    import fitz

    pdf = fitz.open(path)
    try:
        for page_number, page in enumerate(pdf, start=1):
            document.add_paragraph(f'第 {page_number} 页')
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image_stream = BytesIO(pixmap.tobytes('png'))
            _add_image_to_document(document, image_stream)
    finally:
        pdf.close()


def _build_merged_word_file(field, uploaded_records):
    from io import BytesIO
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    image_exts = {'.jpg', '.jpeg', '.png', '.bmp', '.gif', '.webp'}
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

    for style_name in ['Normal', 'Title', 'Heading 1', 'Heading 2']:
        style = document.styles[style_name]
        style.font.name = '微软雅黑'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    document.styles['Normal'].font.size = Pt(10.5)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title_run = title.add_run(field.label)
    title_run.bold = True
    title_run.font.name = '微软雅黑'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    title_run.font.size = Pt(16)

    intro = document.add_paragraph('以下内容由系统根据本字段一次上传的多个文件自动合并生成。')
    intro.paragraph_format.space_before = Pt(0)
    intro.paragraph_format.space_after = Pt(10)

    for index, uploaded in enumerate(uploaded_records, start=1):
        original_name = uploaded.original_name or os.path.basename(uploaded.file.name)
        ext = os.path.splitext(original_name)[1].lower()
        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(4)
        heading.paragraph_format.space_after = Pt(6)
        heading_run = heading.add_run(f'{index}. {original_name}')
        heading_run.bold = True
        heading_run.font.name = '微软雅黑'
        heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        heading_run.font.size = Pt(12)
        try:
            # path() 在 S3 模式下会自动下载到临时文件并返回路径
            path = uploaded.file.path
            if ext in image_exts:
                _add_image_to_document(document, path)
            elif ext == '.pdf':
                _render_pdf_pages_to_document(document, path)
            elif ext == '.docx':
                _append_docx_text(document, path)
            elif ext == '.doc':
                document.add_paragraph('旧版 .doc 文件无法在当前环境中直接解析，原文件已随提交一并保存。')
            else:
                document.add_paragraph('该文件类型不支持合并进 Word，原文件已随提交一并保存。')
        except Exception as exc:
            document.add_paragraph(f'合并该文件时遇到问题：{exc}')
        finally:
            # S3 模式下 path() 会下载到 /tmp，用完立即释放
            try:
                from .storage_backends import cleanup_temp_files
                cleanup_temp_files()
            except Exception:
                pass

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return ContentFile(buffer.getvalue(), name=f'{_safe_filename_part(field.label)}-合并文档.docx')


def _create_uploaded_records(submission, field, uploaded_files):
    created_uploads = []
    total = len(uploaded_files)
    for index, uploaded in enumerate(uploaded_files, start=1):
        source_name = uploaded.name
        renamed_name = _renamed_upload_name(field, uploaded, index=index, total=total)
        uploaded.name = renamed_name
        created_uploads.append(FormUploadedFile.objects.create(
            submission=submission,
            field=field,
            file=uploaded,
            original_name=renamed_name,
            source_name=source_name,
        ))
    return created_uploads


def _refresh_generated_merge_file(submission, field):
    for uploaded in submission.uploaded_files.filter(field=field, is_generated=True):
        _delete_uploaded_record(uploaded)
    source_uploads = list(
        submission.uploaded_files.filter(field=field, is_generated=False).exclude(review_status='rejected').order_by('uploaded_at', 'id')
    )
    if field.merge_files_to_word() and len(source_uploads) > 1:
        merged_file = _build_merged_word_file(field, source_uploads)
        FormUploadedFile.objects.create(
            submission=submission,
            field=field,
            file=merged_file,
            original_name=merged_file.name,
            is_generated=True,
            review_status='approved',
        )


def _ensure_generated_merge_files(submission):
    file_field_ids = submission.uploaded_files.filter(is_generated=False).values_list('field_id', flat=True).distinct()
    fields = FormField.objects.filter(id__in=file_field_ids)
    for field in fields:
        has_generated = submission.uploaded_files.filter(field=field, is_generated=True).exists()
        source_count = submission.uploaded_files.filter(field=field, is_generated=False).exclude(review_status='rejected').count()
        if field.merge_files_to_word() and source_count > 1 and not has_generated:
            _refresh_generated_merge_file(submission, field)


def _refresh_submission_merge_files(submission):
    file_field_ids = submission.uploaded_files.filter(is_generated=False).values_list('field_id', flat=True).distinct()
    for field in FormField.objects.filter(id__in=file_field_ids):
        if field.merge_files_to_word():
            _refresh_generated_merge_file(submission, field)


def _save_dynamic_submission(channel, club, user, fields, cleaned, upload_map, cycle=None):
    previous = FormSubmission.objects.filter(channel=channel, club=club, submitter=user)
    if cycle:
        previous = previous.filter(cycle=cycle)
    previous_count = previous.count()
    submission = FormSubmission.objects.create(
        channel=channel,
        club=club,
        submitter=user,
        cycle=cycle,
        resubmission_count=previous_count + 1,
    )
    for field in fields:
        if field.field_type == 'file':
            uploaded_files = upload_map.get(field.id, [])
            _create_uploaded_records(submission, field, uploaded_files)
            _refresh_generated_merge_file(submission, field)
            continue
        value = cleaned.get(field.id, [] if field.field_type == 'checkbox' else '')
        if field.field_type == 'checkbox':
            FormFieldValue.objects.create(submission=submission, field=field, value_json=value)
        else:
            FormFieldValue.objects.create(submission=submission, field=field, value_text=str(value))
    return submission


def _apply_builtin_action(submission):
    apply_business_action(submission)


def _office_preview_url(request, uploaded):
    if not uploaded.file:
        return ''
    filename = uploaded.original_name or uploaded.file.name
    ext = os.path.splitext(filename)[1].lower()
    if ext not in {'.doc', '.docx', '.ppt', '.pptx'}:
        return ''
    # 走存储抽象层：S3 模式返回 S3/CDN 直链（不经本站代理），
    # 本地模式返回相对 MEDIA_URL，再由 build_absolute_uri 补全为绝对 URL。
    from .storage_backends import ClubStorage
    storage = ClubStorage()
    try:
        direct_url = storage.get_public_url(uploaded.file.name)
    except Exception:
        # 兜底：用 Django 默认 storage.url
        direct_url = uploaded.file.url

    # 如果是本地存储的相对 URL，再用 request 补全为绝对 URL（让 Microsoft 服务器能访问到）
    if direct_url.startswith('/'):
        absolute_file_url = request.build_absolute_uri(direct_url) if request else direct_url
    else:
        # S3 直链已经是绝对 URL，不再包装
        absolute_file_url = direct_url
    encoded_file_url = urllib.parse.quote(absolute_file_url, safe='')
    return f'https://view.officeapps.live.com/op/embed.aspx?src={encoded_file_url}'


def _submission_context(submission, request=None):
    _ensure_generated_merge_files(submission)
    values = {value.field_id: value for value in submission.values.select_related('field')}
    files_by_field = defaultdict(list)
    for uploaded in submission.uploaded_files.filter(is_generated=False).select_related('field'):
        uploaded.office_preview_url = _office_preview_url(request, uploaded)
        files_by_field[uploaded.field_id].append(uploaded)
    rows = []
    for field in submission.channel.fields.filter(is_active=True).order_by('order', 'id'):
        value_obj = values.get(field.id)
        rows.append({
            'field': field,
            'value_obj': value_obj,
            'value': value_obj.value_json if value_obj and value_obj.value_json not in ({}, []) else (value_obj.value_text if value_obj else ''),
            'files': files_by_field.get(field.id, []),
            'review_status': value_obj.review_status if value_obj else '',
            'review_comment': value_obj.review_comment if value_obj else '',
        })
    return rows


def _submission_has_files(submission):
    return submission.uploaded_files.exclude(file='').exists()


def _submission_review_summary(submission):
    current_reviews = submission.reviews.filter(submission_attempt=submission.resubmission_count).select_related('reviewer', 'reviewer__profile')
    approved_count = current_reviews.filter(status='approved').count()
    rejected_count = current_reviews.filter(status='rejected').count()
    required_count = max(1, int(submission.channel.required_approval_count or 1))
    return {
        'current_reviews': current_reviews.order_by('-reviewed_at'),
        'all_reviews': submission.reviews.select_related('reviewer', 'reviewer__profile').order_by('-reviewed_at'),
        'approved_count': approved_count,
        'rejected_count': rejected_count,
        'required_approval_count': required_count,
        'remaining_approval_count': max(0, required_count - approved_count),
        'approval_progress': f'{approved_count}/{required_count}',
    }


def _submission_generated_merge_files(submission):
    _refresh_submission_merge_files(submission)
    return list(submission.uploaded_files.filter(is_generated=True).exclude(file='').select_related('field'))


def _show_word_downloads(submission):
    return submission.status == 'approved'


def _show_zip_download(submission):
    return bool(submission.channel.show_zip_download) and _submission_has_files(submission)


def _submission_download_context(submission):
    show_word_downloads = _show_word_downloads(submission)
    show_zip_download = _show_zip_download(submission)
    return {
        'has_files': _submission_has_files(submission),
        'show_word_downloads': show_word_downloads,
        'show_zip_download': show_zip_download,
        'generated_merge_files': _submission_generated_merge_files(submission) if show_word_downloads else [],
    }


def _submission_common_context(submission, request=None):
    context = {
        'rows': _submission_context(submission, request=request),
    }
    context.update(_submission_review_summary(submission))
    context.update(_submission_download_context(submission))
    return context


def _rejected_field_queryset(submission):
    rejected_value_ids = submission.values.filter(review_status='rejected').values_list('field_id', flat=True)
    rejected_file_ids = submission.uploaded_files.filter(review_status='rejected', is_generated=False).values_list('field_id', flat=True)
    return submission.channel.fields.filter(id__in=set(rejected_value_ids) | set(rejected_file_ids), is_active=True).order_by('order', 'id')


def _form_field_items(fields, submission=None):
    value_map = {}
    selected_map = {}
    comment_map = {}
    files_by_field = defaultdict(list)
    if submission:
        for value in submission.values.select_related('field'):
            raw = value.value_json if value.value_json not in ({}, []) else value.value_text
            value_map[value.field_id] = raw
            selected_map[value.field_id] = raw if isinstance(raw, list) else [raw]
            comment_map[value.field_id] = value.review_comment
        for uploaded in submission.uploaded_files.filter(is_generated=False).select_related('field'):
            files_by_field[uploaded.field_id].append(uploaded)
    items = []
    for field in fields:
        value = value_map.get(field.id, [] if field.field_type == 'checkbox' else '')
        items.append({
            'field': field,
            'value': value,
            'selected_options': selected_map.get(field.id, []),
            'review_comment': comment_map.get(field.id, ''),
            'files': files_by_field.get(field.id, []),
        })
    return items


@login_required(login_url=settings.LOGIN_URL)
def submit_dynamic_form(request, channel_slug, club_id):
    channel = _get_channel(channel_slug)
    club = get_object_or_404(Club, pk=club_id)
    if not _is_president(request.user) or not _is_president_of_club(request.user, club):
        messages.error(request, '只有该社团现任社长可以提交')
        return redirect('clubs:user_dashboard')

    can_submit, cycle, policy_error = _can_create_submission(channel, club, request.user)
    if not can_submit:
        messages.error(request, policy_error)
        return redirect('clubs:user_dashboard')

    if request.method == 'POST':
        fields, cleaned, upload_map, errors = _validate_dynamic_submission(channel, request.POST, request.FILES)
        if errors:
            for error in errors:
                messages.error(request, error)
        else:
            submission = _save_dynamic_submission(channel, club, request.user, fields, cleaned, upload_map, cycle=cycle)
            messages.success(request, f'{channel.name} 已提交，等待审核')
            return redirect('clubs:approval_detail', item_type=channel.slug, submission_key=submission.public_id)
    else:
        fields = channel.fields.filter(is_active=True).order_by('order', 'id')

    return render(request, 'clubs/user/dynamic_form_submit.html', {
        'channel': channel,
        'club': club,
        'fields': fields,
        'field_items': _form_field_items(fields),
        'cycle': cycle,
        'revision_mode': False,
    })


@login_required(login_url=settings.LOGIN_URL)
def revise_dynamic_submission(request, submission_key):
    submission = _get_submission_or_404(submission_key)
    if not _is_president(request.user) or not _is_president_of_club(request.user, submission.club):
        messages.error(request, '只有该社团现任社长可以修改补交')
        return redirect('clubs:user_dashboard')
    if submission.status != 'rejected':
        messages.error(request, '只有被打回的请求可以修改补交')
        return redirect('clubs:approval_detail', item_type=submission.channel.slug, submission_key=submission.public_id)

    fields = _rejected_field_queryset(submission)
    if not fields.exists():
        fields = submission.channel.fields.filter(is_active=True).order_by('order', 'id')
    force_required_ids = set(fields.values_list('id', flat=True))

    if request.method == 'POST':
        fields, cleaned, upload_map, errors = _validate_dynamic_submission(
            submission.channel,
            request.POST,
            request.FILES,
            fields=fields,
            force_required_field_ids=force_required_ids,
        )
        if errors:
            for error in errors:
                messages.error(request, error)
        else:
            with transaction.atomic():
                for field in fields:
                    if field.field_type == 'file':
                        for uploaded in submission.uploaded_files.filter(field=field, review_status='rejected', is_generated=False):
                            _delete_uploaded_record(uploaded)
                        new_uploads = _create_uploaded_records(submission, field, upload_map.get(field.id, []))
                        for uploaded in new_uploads:
                            uploaded.review_status = 'pending'
                            uploaded.review_comment = ''
                            uploaded.save(update_fields=['review_status', 'review_comment'])
                        _refresh_generated_merge_file(submission, field)
                        continue

                    value = cleaned.get(field.id, [] if field.field_type == 'checkbox' else '')
                    defaults = {'review_status': 'pending', 'review_comment': ''}
                    if field.field_type == 'checkbox':
                        defaults['value_json'] = value
                        defaults['value_text'] = ''
                    else:
                        defaults['value_text'] = str(value)
                        defaults['value_json'] = {}
                    FormFieldValue.objects.update_or_create(
                        submission=submission,
                        field=field,
                        defaults=defaults,
                    )
                submission.status = 'pending'
                submission.review_comment = ''
                submission.reviewer = None
                submission.reviewed_at = None
                submission.submitted_at = timezone.now()
                submission.resubmission_count += 1
                submission.is_read = False
                submission.save(update_fields=[
                    'status',
                    'review_comment',
                    'reviewer',
                    'reviewed_at',
                    'submitted_at',
                    'resubmission_count',
                    'is_read',
                ])
            messages.success(request, '已补交被打回的内容，等待重新审核')
            return redirect('clubs:approval_detail', item_type=submission.channel.slug, submission_key=submission.public_id)

    return render(request, 'clubs/user/dynamic_form_submit.html', {
        'channel': submission.channel,
        'club': submission.club,
        'fields': fields,
        'field_items': _form_field_items(fields, submission=submission),
        'cycle': submission.cycle,
        'submission': submission,
        'revision_mode': True,
    })


def _redirect_builtin_submit(request, club_id, slug):
    return redirect('clubs:submit_dynamic_form', channel_slug=slug, club_id=club_id)














@login_required(login_url=settings.LOGIN_URL)
def approval_center_tabs(request, tab='all'):
    if not _is_president(request.user):
        messages.error(request, '仅社长可以访问审批记录')
        return redirect('clubs:index')
    clubs = Club.objects.filter(officers__user_profile__user=request.user, officers__position='president', officers__is_current=True)
    items = FormSubmission.objects.filter(club__in=clubs).select_related('channel', 'club', 'reviewer').order_by('-submitted_at')
    if tab and tab != 'all':
        items = items.filter(channel__slug=tab.replace('_', '-'))
    active_items = items.filter(status__in=['pending', 'rejected'])
    completed_items = items.exclude(status__in=['pending', 'rejected'])
    channels = list(_active_channels())
    grouped_channels = []
    for channel in channels:
        channel_items = items.filter(channel=channel)
        grouped_channels.append({
            'channel': channel,
            'active_items': channel_items.filter(status__in=['pending', 'rejected']),
            'completed_items': channel_items.exclude(status__in=['pending', 'rejected']),
            'total_count': channel_items.count(),
        })
    return render(request, 'clubs/user/dynamic_approval_center.html', {
        'items': items,
        'active_items': active_items,
        'completed_items': completed_items,
        'channels': channels,
        'grouped_channels': grouped_channels,
        'current_tab': tab.replace('_', '-'),
    })




@login_required(login_url=settings.LOGIN_URL)
def approval_detail(request, item_type, submission_key):
    submission = _get_submission_or_404(submission_key)
    if _is_president(request.user) and not _is_president_of_club(request.user, submission.club):
        messages.error(request, '无权查看此提交')
        return redirect('clubs:user_dashboard')
    if not (_is_president(request.user) or is_staff_or_admin(request.user)):
        messages.error(request, '无权查看此提交')
        return redirect('clubs:index')
    context = {
        'item': submission,
        'submission': submission,
        'item_type': submission.channel.slug,
    }
    context.update(_submission_common_context(submission, request=request))
    return render(request, 'clubs/user/dynamic_approval_detail.html', context)


@login_required(login_url=settings.LOGIN_URL)
def staff_audit_center(request, tab='all'):
    if not is_staff_or_admin(request.user):
        messages.error(request, '仅干事和管理员可以访问审核中心')
        return redirect('clubs:index')
    slug = tab.replace('_', '-')
    channels = list(_active_channels())
    current_channel = None
    if slug != 'all':
        current_channel = FormChannel.objects.filter(slug=slug).first()
    qs = FormSubmission.objects.select_related('channel', 'club', 'submitter', 'reviewer').prefetch_related('reviews').order_by('-submitted_at')
    if current_channel:
        qs = qs.filter(channel=current_channel)
    pending_items = qs.filter(status='pending')
    completed_items = qs.exclude(status='pending')
    return render(request, 'clubs/staff/dynamic_audit_center.html', {
        'channels': channels,
        'current_channel': current_channel,
        'current_tab': slug,
        'pending_items': pending_items,
        'completed_items': completed_items,
        'is_admin': _is_admin(request.user),
    })




def _mark_submission_approved(submission, reviewer, comment):
    if FormSubmissionReview.objects.filter(
        submission=submission,
        reviewer=reviewer,
        submission_attempt=submission.resubmission_count,
    ).exists():
        raise BusinessActionError('您已经审核过本次提交，不能重复审核')

    FormSubmissionReview.objects.create(
        submission=submission,
        reviewer=reviewer,
        status='approved',
        comment=comment,
        submission_attempt=submission.resubmission_count,
    )
    approved_count = submission.approved_review_count()
    required_count = submission.required_approval_count

    submission.reviewer = reviewer
    submission.reviewed_at = timezone.now()
    submission.is_read = False
    submission.review_comment = comment
    if approved_count >= required_count:
        submission.values.update(review_status='approved', review_comment='')
        submission.uploaded_files.update(review_status='approved', review_comment='')
        submission.status = 'approved'
        submission.save()
        _apply_builtin_action(submission)
    else:
        submission.status = 'pending'
        submission.save()


def _mark_submission_rejected(submission, reviewer, comment, post):
    if FormSubmissionReview.objects.filter(
        submission=submission,
        reviewer=reviewer,
        submission_attempt=submission.resubmission_count,
    ).exists():
        raise BusinessActionError('您已经审核过本次提交，不能重复审核')

    FormSubmissionReview.objects.create(
        submission=submission,
        reviewer=reviewer,
        status='rejected',
        comment=comment,
        submission_attempt=submission.resubmission_count,
    )

    rejected_value_ids = {
        int(key.rsplit('_', 1)[1])
        for key in post
        if key.startswith('reject_value_') and key.rsplit('_', 1)[1].isdigit()
    }
    rejected_file_ids = {
        int(key.rsplit('_', 1)[1])
        for key in post
        if key.startswith('reject_file_') and key.rsplit('_', 1)[1].isdigit()
    }

    if not rejected_value_ids and not rejected_file_ids:
        rejected_value_ids = set(submission.values.values_list('id', flat=True))
        rejected_file_ids = set(submission.uploaded_files.filter(is_generated=False).values_list('id', flat=True))

    for value in submission.values.select_related('field'):
        is_rejected = value.id in rejected_value_ids
        value.review_status = 'rejected' if is_rejected else 'approved'
        value.review_comment = post.get(f'comment_value_{value.id}', '').strip() if is_rejected else ''
        if is_rejected and not value.review_comment:
            value.review_comment = comment
        value.save(update_fields=['review_status', 'review_comment', 'updated_at'])

    for uploaded in submission.uploaded_files.select_related('field'):
        if uploaded.is_generated:
            uploaded.review_status = 'approved'
            uploaded.review_comment = ''
        else:
            is_rejected = uploaded.id in rejected_file_ids
            uploaded.review_status = 'rejected' if is_rejected else 'approved'
            uploaded.review_comment = post.get(f'comment_file_{uploaded.id}', '').strip() if is_rejected else ''
            if is_rejected and not uploaded.review_comment:
                uploaded.review_comment = comment
        uploaded.save(update_fields=['review_status', 'review_comment'])

    for field_id in submission.uploaded_files.filter(is_generated=False).values_list('field_id', flat=True).distinct():
        field = FormField.objects.filter(id=field_id).first()
        if field:
            _refresh_generated_merge_file(submission, field)

    submission.status = 'rejected'
    submission.review_comment = comment
    submission.reviewer = reviewer
    submission.reviewed_at = timezone.now()
    submission.is_read = False
    submission.save()


@login_required(login_url=settings.LOGIN_URL)
def staff_review_form_submission(request, submission_key):
    if not is_staff_or_admin(request.user):
        messages.error(request, '仅干事和管理员可以审核')
        return redirect('clubs:index')
    submission = _get_submission_or_404(submission_key)
    if request.method == 'POST':
        decision = request.POST.get('decision')
        comment = request.POST.get('comment', '').strip()
        if decision not in ['approved', 'rejected']:
            messages.error(request, '请选择有效审核结果')
        else:
            try:
                with transaction.atomic():
                    if decision == 'approved':
                        _mark_submission_approved(submission, request.user, comment)
                    else:
                        _mark_submission_rejected(submission, request.user, comment, request.POST)
            except BusinessActionError as exc:
                messages.error(request, str(exc))
            else:
                if decision == 'approved':
                    approved_count = submission.approved_review_count()
                    required_count = submission.required_approval_count
                    if submission.status == 'approved':
                        messages.success(request, f'审核结果已保存，已达到 {approved_count}/{required_count} 次通过')
                    else:
                        messages.success(request, f'已记录本次通过，当前通过进度 {approved_count}/{required_count}')
                else:
                    messages.success(request, '审核结果已保存，提交已打回')
                if decision == 'approved' and (_show_word_downloads(submission) or _show_zip_download(submission)):
                    return redirect('clubs:staff_review_form_submission', submission_key=submission.public_id)
                return redirect('clubs:staff_audit_center', tab=submission.channel.slug)
    context = {
        'submission': submission,
    }
    context.update(_submission_common_context(submission, request=request))
    context['user_has_reviewed_current_attempt'] = submission.reviews.filter(
        reviewer=request.user,
        submission_attempt=submission.resubmission_count,
    ).exists()
    return render(request, 'clubs/staff/dynamic_submission_review.html', context)


@login_required(login_url=settings.LOGIN_URL)
@require_POST
def delete_audit_request(request, tab, item_key):
    if not _is_admin(request.user):
        messages.error(request, '仅管理员可以删除审核请求')
        return redirect('clubs:staff_audit_center', tab=tab)
    submission = get_object_or_404(FormSubmission, public_id=item_key)
    slug = submission.channel.slug
    for uploaded in submission.uploaded_files.all():
        if uploaded.file:
            uploaded.file.delete(save=False)
    submission.delete()
    messages.success(request, '审核请求已删除')
    return redirect('clubs:staff_audit_center', tab=slug)


@login_required(login_url=settings.LOGIN_URL)
def manage_form_channels(request, channel_id=None):
    if not _is_admin(request.user):
        messages.error(request, '仅管理员可以管理提交通道')
        return redirect('clubs:index')
    channels = FormChannel.objects.prefetch_related('fields', 'cycles').order_by('order', 'id')
    is_creating = request.GET.get('new') == '1'
    current = None if is_creating else (get_object_or_404(FormChannel, pk=channel_id) if channel_id else channels.first())
    missing_fields = missing_required_field_keys(current) if current else []
    locked_field_keys = locked_business_field_keys(current) if current else set()
    return render(request, 'clubs/admin/form_channels.html', {
        'channels': channels,
        'current': current,
        'is_creating': is_creating or current is None,
        'field_types': FormField.FIELD_TYPE_CHOICES,
        'builtin_actions': FormChannel.BUILTIN_ACTION_CHOICES,
        'submission_policies': FormChannel.SUBMISSION_POLICY_CHOICES,
        'cycle_types': FormChannel.CYCLE_TYPE_CHOICES,
        'missing_required_fields': missing_fields,
        'locked_field_keys': locked_field_keys,
    })


@login_required(login_url=settings.LOGIN_URL)
@require_POST
def save_form_channel(request, channel_id=None):
    if not _is_admin(request.user):
        messages.error(request, '仅管理员可以管理提交通道')
        return redirect('clubs:index')
    channel = get_object_or_404(FormChannel, pk=channel_id) if channel_id else FormChannel()
    channel.name = request.POST.get('name', '').strip()
    channel.slug = request.POST.get('slug', '').strip()
    channel.icon = request.POST.get('icon', 'description').strip() or 'description'
    channel.description = request.POST.get('description', '').strip()
    channel.builtin_action = request.POST.get('builtin_action', 'none')
    channel.submission_policy = request.POST.get('submission_policy', 'repeatable')
    channel.cycle_type = request.POST.get('cycle_type', 'none')
    try:
        channel.required_approval_count = min(9, max(1, int(request.POST.get('required_approval_count', '1') or 1)))
    except ValueError:
        channel.required_approval_count = 1
    if channel.cycle_type != 'none':
        channel.submission_policy = 'once_per_cycle'
    elif channel.submission_policy == 'once_per_cycle':
        channel.submission_policy = 'repeatable'
    channel.is_active = request.POST.get('is_active') == 'on'
    channel.is_builtin = request.POST.get('is_builtin') == 'on'
    channel.show_zip_download = request.POST.get('show_zip_download') == 'on'
    channel.show_unsubmitted_status = request.POST.get('show_unsubmitted_status') == 'on'
    channel.allow_staff_toggle = request.POST.get('allow_staff_toggle') == 'on'
    try:
        channel.order = int(request.POST.get('order', '0') or 0)
    except ValueError:
        channel.order = 0
    if not channel.name or not channel.slug:
        messages.error(request, '通道名称和标识不能为空')
    else:
        try:
            channel.save()
            messages.success(request, '通道已保存')
        except IntegrityError:
            messages.error(request, '通道标识已存在，请换一个 slug')
    return redirect('clubs:manage_form_channels_detail', channel_id=channel.id) if channel.id else redirect('clubs:manage_form_channels')


@login_required(login_url=settings.LOGIN_URL)
@require_POST
def delete_form_channel(request, channel_id):
    if not _is_admin(request.user):
        messages.error(request, '仅管理员可以删除提交通道')
        return redirect('clubs:index')
    channel = get_object_or_404(FormChannel, pk=channel_id)
    channel.delete()
    messages.success(request, '通道已删除')
    return redirect('clubs:manage_form_channels')


@login_required(login_url=settings.LOGIN_URL)
@require_POST
def save_form_field(request, channel_id, field_id=None):
    if not _is_admin(request.user):
        messages.error(request, '仅管理员可以管理字段')
        return redirect('clubs:index')
    channel = get_object_or_404(FormChannel, pk=channel_id)
    field = get_object_or_404(FormField, pk=field_id, channel=channel) if field_id else FormField(channel=channel)
    if field_id and is_locked_business_field(field):
        messages.error(request, '这个字段已绑定内置业务逻辑，不能修改')
        return redirect('clubs:manage_form_channels_detail', channel_id=channel.id)
    field.label = request.POST.get('label', '').strip()
    field.field_key = request.POST.get('field_key', '').strip()
    field.field_type = request.POST.get('field_type', 'text')
    field.required = request.POST.get('required') == 'on'
    field.help_text = request.POST.get('help_text', '').strip()
    field.placeholder = request.POST.get('placeholder', '').strip()
    field.options = _parse_options(request.POST.get('options', ''))
    field.validation = _parse_validation(field.field_type, request.POST)
    if (
        field.field_type == 'file'
        and request.POST.get('merge_to_word') == 'on'
        and not field.validation.get('merge_to_word')
    ):
        messages.warning(request, '合并 Word 只支持开启多文件后，且允许后缀全部为图片、PDF 或 Word（.docx）的文件字段')
    field.is_active = request.POST.get('is_active') == 'on'
    try:
        field.order = int(request.POST.get('order', '0') or 0)
    except ValueError:
        field.order = 0
    if 'example_file' in request.FILES:
        field.example_file = request.FILES['example_file']
    if request.POST.get('clear_example') == 'on':
        field.example_file = None
    if not field.label or not field.field_key:
        messages.error(request, '字段名称和标识不能为空')
    else:
        try:
            field.save()
            messages.success(request, '字段已保存')
        except IntegrityError:
            messages.error(request, '字段标识已存在，请换一个 field_key')
    return redirect('clubs:manage_form_channels_detail', channel_id=channel.id)


@login_required(login_url=settings.LOGIN_URL)
@require_POST
def delete_form_field(request, channel_id, field_id):
    if not _is_admin(request.user):
        messages.error(request, '仅管理员可以删除字段')
        return redirect('clubs:index')
    field = get_object_or_404(FormField, pk=field_id, channel_id=channel_id)
    if is_locked_business_field(field):
        messages.error(request, '这个字段已绑定内置业务逻辑，不能删除')
        return redirect('clubs:manage_form_channels_detail', channel_id=channel_id)
    field.delete()
    messages.success(request, '字段已删除')
    return redirect('clubs:manage_form_channels_detail', channel_id=channel_id)


@login_required(login_url=settings.LOGIN_URL)
@require_POST
def create_form_channel_cycle(request, channel_id):
    if not _is_admin(request.user):
        messages.error(request, '仅管理员可以管理周期')
        return redirect('clubs:index')
    channel = get_object_or_404(FormChannel, pk=channel_id)
    name = request.POST.get('name', '').strip()
    cycle = create_form_cycle(channel, name=name, user=request.user)
    messages.success(request, f'周期已创建：{cycle.name}')
    return redirect('clubs:manage_form_channels_detail', channel_id=channel_id)


@login_required(login_url=settings.LOGIN_URL)
@require_POST
def close_form_channel_cycle(request, channel_id, cycle_id):
    if not _is_admin(request.user):
        messages.error(request, '仅管理员可以管理周期')
        return redirect('clubs:index')
    cycle = get_object_or_404(FormCycle, pk=cycle_id, channel_id=channel_id)
    cycle.is_active = False
    cycle.ends_at = timezone.now()
    cycle.save(update_fields=['is_active', 'ends_at'])
    messages.success(request, f'周期已关闭：{cycle.name}')
    return redirect('clubs:manage_form_channels_detail', channel_id=channel_id)


@login_required(login_url=settings.LOGIN_URL)
def admin_dashboard(request):
    if not _is_admin(request.user):
        messages.error(request, '仅管理员可以访问此页面')
        return redirect('clubs:index')
    total_clubs = Club.objects.count()
    total_users = User.objects.count()
    pending_registrations = FormSubmission.objects.filter(status='pending').count()
    published_announcements = Announcement.objects.filter(status='published').count()
    pending_staff_count = UserProfile.objects.filter(role='staff', status='pending').count()
    announcements = Announcement.objects.all().order_by('-created_at')[:5]
    total_applications = FormSubmission.objects.count()
    pending_all = FormSubmission.objects.filter(status='pending').count()
    channels_count = FormChannel.objects.count()

    today = timezone.localdate()
    date_range = [today - timezone.timedelta(days=offset) for offset in range(13, -1, -1)]
    daily_stats = {
        item.date: item.visits
        for item in DailyStat.objects.filter(date__range=(date_range[0], date_range[-1]))
    }
    visit_dates = [item.isoformat() for item in date_range]
    visit_counts = [daily_stats.get(item, 0) for item in date_range]

    status_rows = FormSubmission.objects.values('channel_id', 'channel__name', 'status').annotate(total=Count('id'))
    status_map = {
        (row['channel_id'], row['status']): row['total']
        for row in status_rows
    }
    chart_channels = list(FormChannel.objects.order_by('order', 'id').values('id', 'name'))
    known_channel_ids = {channel['id'] for channel in chart_channels}
    for row in status_rows:
        if row['channel_id'] not in known_channel_ids:
            chart_channels.append({'id': row['channel_id'], 'name': row['channel__name'] or '未知通道'})
    type_labels = [channel['name'] for channel in chart_channels]
    type_pending = [status_map.get((channel['id'], 'pending'), 0) for channel in chart_channels]
    type_approved = [status_map.get((channel['id'], 'approved'), 0) for channel in chart_channels]
    type_rejected = [status_map.get((channel['id'], 'rejected'), 0) for channel in chart_channels]

    approved_total = FormSubmission.objects.filter(status='approved').count()
    rejected_total = FormSubmission.objects.filter(status='rejected').count()
    decided_total = approved_total + rejected_total
    overall_approval_rate = round(approved_total / decided_total * 100, 1) if decided_total else 0
    overall_rejection_rate = round(rejected_total / decided_total * 100, 1) if decided_total else 0

    return render(request, 'clubs/admin/dashboard.html', {
        'total_clubs': total_clubs,
        'total_users': total_users,
        'pending_registrations': pending_registrations,
        'published_announcements': published_announcements,
        'pending_staff_count': pending_staff_count,
        'presidents_count': UserProfile.objects.filter(role='president').count(),
        'staff_count': UserProfile.objects.filter(role='staff').count(),
        'admins_count': UserProfile.objects.filter(role='admin').count(),
        'members_count': UserProfile.objects.filter(role='member').count(),
        'announcements': announcements,
        'total_applications': total_applications,
        'pending_all': pending_all,
        'channels_count': channels_count,
        'overall_approval_rate': overall_approval_rate,
        'overall_rejection_rate': overall_rejection_rate,
        'visit_dates': visit_dates,
        'visit_counts': visit_counts,
        'visit_dates_json': json.dumps(visit_dates, ensure_ascii=False),
        'visit_counts_json': json.dumps(visit_counts, ensure_ascii=False),
        'total_visits_14d': sum(visit_counts),
        'type_labels_json': json.dumps(type_labels, ensure_ascii=False),
        'type_pending_json': json.dumps(type_pending, ensure_ascii=False),
        'type_approved_json': json.dumps(type_approved, ensure_ascii=False),
        'type_rejected_json': json.dumps(type_rejected, ensure_ascii=False),
        'redis_info': None,
    })



@login_required
def public_activities(request):
    role = getattr(getattr(request.user, 'profile', None), 'role', None)
    if role not in ['staff', 'admin', 'president', 'member']:
        messages.error(request, '您没有权限访问此页面')
        return redirect('clubs:user_dashboard')

    qs = PublishedActivity.objects.select_related('club', 'source_submission').order_by('-published_at')
    if role == 'president':
        qs = qs.filter(club_id__in=_get_president_club_ids(request.user))
    elif role == 'member':
        member_club_ids = list(ClubMember.objects.filter(user_profile__user=request.user, status='active').values_list('club_id', flat=True))
        qs = qs.filter(Q(club_id__in=member_club_ids) | Q(is_public=True)).distinct()

    activities = list(qs)
    search_query = request.GET.get('search', '').strip()
    if search_query:
        activities = [item for item in activities if search_query in item.activity_name or search_query in item.club.name]
    activity_type_filter = request.GET.get('activity_type', '').strip()
    if activity_type_filter:
        activities = [item for item in activities if item.activity_type == activity_type_filter]
    club_filter = request.GET.get('club', '').strip()
    if club_filter:
        activities = [item for item in activities if club_filter in item.club.name]
    date_filter = request.GET.get('date', '').strip()
    if date_filter:
        activities = [item for item in activities if str(item.activity_date) == date_filter]

    activities_by_type = defaultdict(list)
    for activity in activities:
        activities_by_type[activity.get_activity_type_display()].append(activity)

    registered_ids = set()
    if role == 'member':
        registered_ids = set(ActivityRegistration.objects.filter(user_profile__user=request.user).values_list('activity_id', flat=True))
    return render(request, 'clubs/public_activities.html', {
        'approved_activities': activities,
        'activities_by_type': dict(activities_by_type),
        'all_clubs': Club.objects.all().order_by('name'),
        'activity_type_choices': PublishedActivity.ACTIVITY_TYPE_CHOICES,
        'club_filter': club_filter,
        'activity_type_filter': activity_type_filter,
        'date_filter': date_filter,
        'search_query': search_query,
        'user_registered_ids': registered_ids,
    })


@login_required
@require_POST
def register_activity(request, activity_id):
    if getattr(request.user.profile, 'role', None) != 'member':
        return JsonResponse({'success': False, 'error': '仅社员可以报名活动'}, status=403)
    activity = get_object_or_404(PublishedActivity, pk=activity_id)
    if not activity.is_public and not ClubMember.objects.filter(user_profile=request.user.profile, club=activity.club, status='active').exists():
        return JsonResponse({'success': False, 'error': '您不是该社团成员，无法报名'}, status=403)
    _, created = ActivityRegistration.objects.get_or_create(activity=activity, user_profile=request.user.profile)
    if created:
        return JsonResponse({'success': True, 'registered': True})
    return JsonResponse({'success': False, 'error': '您已报名该活动'})


@login_required
@require_POST
def unregister_activity(request, activity_id):
    if getattr(request.user.profile, 'role', None) != 'member':
        return JsonResponse({'success': False, 'error': '仅社员可以取消报名'}, status=403)
    deleted, _ = ActivityRegistration.objects.filter(activity_id=activity_id, user_profile=request.user.profile).delete()
    if deleted:
        return JsonResponse({'success': True, 'registered': False})
    return JsonResponse({'success': False, 'error': '您尚未报名该活动'})


@login_required(login_url=settings.LOGIN_URL)
@require_http_methods(['GET'])
def notification_counts(request):
    role = getattr(getattr(request.user, 'profile', None), 'role', '')
    audit_counts = {}
    approval_counts = {}
    if role in ['staff', 'admin'] or request.user.is_superuser:
        for row in FormSubmission.objects.filter(status='pending').values('channel__slug').annotate(count=Count('id')):
            audit_counts[row['channel__slug']] = row['count']
    if role == 'president':
        club_ids = _get_president_club_ids(request.user)
        for row in FormSubmission.objects.filter(club_id__in=club_ids, status__in=['pending', 'rejected']).values('channel__slug').annotate(count=Count('id')):
            approval_counts[row['channel__slug']] = row['count']
    return JsonResponse({
        'role': role,
        'audit_counts': audit_counts,
        'approval_counts': {**approval_counts, 'total': sum(approval_counts.values())},
        'audit_total': sum(audit_counts.values()),
    })




@login_required(login_url=settings.LOGIN_URL)
@require_POST
def cancel_submission(request, submission_key):
    submission = get_object_or_404(FormSubmission, public_id=submission_key, status='pending')
    if not _is_president_of_club(request.user, submission.club):
        messages.error(request, '无权取消此提交')
        return redirect('clubs:user_dashboard')
    for uploaded in submission.uploaded_files.all():
        if uploaded.file:
            uploaded.file.delete(save=False)
    submission.delete()
    messages.success(request, '提交已取消')
    return redirect('clubs:approval_center', tab='all')
